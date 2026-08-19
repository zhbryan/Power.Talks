#!/usr/bin/env python3
"""
Shared helper: extract the CHECKED "Reason for Revision" option(s) from an
ERCOT market-rule -01 filing (.docx). Used by every profile_ercot_<cat>.py.

The reason options are ActiveX Forms.TextBox controls; the sponsor types an "X"
into the chosen one(s). python-docx cannot see that "X" (it is not text) and
Word will not instantiate the controls headlessly. We read the state from each
control's embedded OLE stream: a control is CHECKED iff its `contents` stream
carries a stored 1-character value, which the MorphData format records as the
Value-count DWORD 0x80000001 (little-endian b"\\x01\\x00\\x00\\x80"). Empty
(unchecked) controls have no Value property.

This keeps only the box(es) the sponsor actually marked, rather than listing
every option. Works across both ERCOT reason templates (older "Addresses
current operational issues / Meets Strategic goals / ..." and newer "Strategic
Plan Objective 1-3 / ...") because it reads the paragraph text next to each
checked control, and across all six categories (NPRR/COPMGRR/PGRR/SCR/NOGRR/
RMGRR), which share the same form. Legacy .doc filings expose no readable
controls -> [].

Newer template variants render each checkbox as an embedded image (a WMF
glyph) instead of an ActiveX control -- either DrawingML (<a:blip>) or legacy
VML (<v:imagedata>). Every option carries the same empty-box image except the
one the sponsor marked, which carries a different (checked) image. When no
ActiveX controls are present we fall back to comparing those per-option images:
the box image that repeats across most options is the unchecked state, so any
option whose box image differs is the checked one.
"""

import hashlib
import io
import re
import zipfile
from collections import Counter

import olefile
from docx import Document
from docx.oxml.ns import qn

_REASON_NOTE_PAT = re.compile(r'please select', re.IGNORECASE)


def _reason_ctrl_checked(z, relmap, rid):
    tgt = relmap.get(rid)
    if not tgt:
        return False
    try:
        ole = olefile.OleFileIO(io.BytesIO(z.read('word/' + tgt.replace('.xml', '.bin'))))
        data = ole.openstream('contents').read()
        ole.close()
    except Exception:
        return False
    return b'\x01\x00\x00\x80' in data


def _clean_reason(text):
    text = re.sub(r'\s+', ' ', text).replace('�', '-').strip()
    return text.rstrip('.').strip()


def _read_rel_blob(z, relmap, rid):
    """Bytes of the relationship target `rid` (relative to word/), or None."""
    tgt = relmap.get(rid)
    if not tgt:
        return None
    for name in ('word/' + tgt.replace('\\', '/'), tgt.lstrip('/')):
        try:
            return z.read(name)
        except Exception:
            continue
    return None


_REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def _localname(el):
    t = el.tag
    return t.rsplit('}', 1)[-1] if isinstance(t, str) else ''


def _iter_box_images(el):
    """Yield the relationship id of every embedded image under `el`, covering
    both checkbox-image encodings: DrawingML <a:blip r:embed=".."> and legacy
    VML <v:imagedata r:id="..">."""
    for d in el.iter():
        ln = _localname(d)
        if ln == 'blip':
            rid = d.get('{%s}embed' % _REL_NS) or d.get('{%s}link' % _REL_NS)
        elif ln == 'imagedata':
            rid = d.get('{%s}id' % _REL_NS)
        else:
            continue
        if rid:
            yield rid


def _para_box_image_hash(z, relmap, para):
    """md5 of the first embedded checkbox image in the paragraph, or None. In
    the image-checkbox templates that image IS the option's checkbox (either a
    DrawingML a:blip or a legacy VML v:imagedata reference)."""
    for rid in _iter_box_images(para._p):
        blob = _read_rel_blob(z, relmap, rid)
        if blob:
            return hashlib.md5(blob).hexdigest()
    return None


def _image_checked_reasons(z, relmap, cell):
    """Detect the checked option(s) when checkboxes are embedded images rather
    than ActiveX controls. The empty box is the image repeated across most
    options; an option whose box image differs is the one that was marked."""
    opts = []  # (label, image_md5)
    for para in cell.paragraphs:
        t = para.text.strip()
        if not t or _REASON_NOTE_PAT.search(t):
            continue
        h = _para_box_image_hash(z, relmap, para)
        if h:
            opts.append((t, h))
    if len(opts) < 2:
        return []
    modal, _ = Counter(h for _, h in opts).most_common(1)[0]
    return [_clean_reason(t) for t, h in opts if h != modal]


def extract_checked_reasons(docx_path):
    """Return the Reason-for-Revision option(s) the sponsor checked, as the
    paragraph text next to each marked box. Handles both templates: ActiveX
    checkbox controls and embedded-image checkboxes. Returns [] when the file
    is a legacy .doc (no readable markers), nothing is checked, or the reason
    table cannot be located."""
    if not docx_path or not docx_path.lower().endswith('.docx'):
        return []
    try:
        z = zipfile.ZipFile(docx_path)
        rels = z.read('word/_rels/document.xml.rels').decode('utf-8')
    except Exception:
        return []
    # Full relationship map — activeX targets drive the control path, media
    # (image) targets drive the image-checkbox fallback.
    relmap = dict(re.findall(r'Id="(rId\d+)"[^>]*?Target="([^"]+)"', rels))
    try:
        doc = Document(docx_path)
    except Exception:
        return []
    for tbl in doc.tables:
        for row in tbl.rows:
            if 'reason for revision' not in ' '.join(c.text for c in row.cells).lower():
                continue
            # Dedupe merged cells; the options cell bears the controls/images.
            seen, cells = set(), []
            for c in row.cells:
                if id(c._tc) in seen:
                    continue
                seen.add(id(c._tc)); cells.append(c)
            ctrl_cells = [c for c in cells
                          if any(True for _ in c._tc.iter(qn('w:control')))]
            if ctrl_cells:
                cell = max(ctrl_cells, key=lambda c: len(c.text))
                checked = []
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if not t or _REASON_NOTE_PAT.search(t):
                        continue
                    rid = None
                    for ctrl in para._p.iter(qn('w:control')):
                        rid = ctrl.get(qn('r:id'))
                    if rid and _reason_ctrl_checked(z, relmap, rid):
                        checked.append(_clean_reason(t))
                return checked
            # No ActiveX controls: image-checkbox templates (DrawingML a:blip
            # or legacy VML v:imagedata).
            img_cells = [c for c in cells
                         if next(_iter_box_images(c._tc), None) is not None]
            if img_cells:
                cell = max(img_cells, key=lambda c: len(c.text))
                return _image_checked_reasons(z, relmap, cell)
    return []
