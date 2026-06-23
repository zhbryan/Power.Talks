#!/usr/bin/env python3
"""Populate `source_documents` (the "Documents Submitted" list + per-document
summary) on every ERCOT market-rules issue Profile.json.

Implements `Set Paper Trails Document Summary.md`: for each issue under
Documents Database/ERCOT.MKT.RULES/<CAT>/<ISSUE_ID>/, list the submitted files
(excluding .zip/.tmp and the Quick runs JSONs) and write, into the issue's
Profile.json, a `source_documents` array of:

    { "file", "doc_type", "date", "author", "summary", "key_points",
      "download_url" }

The Paper Trails item-rule card lists these (title link + download button) and
the content window renders the clicked entry's summary (see the homepage skill).

Summaries are filename-derived by default (fast, full coverage). Pass --read to
also open each document and extract a couple of substantive sentences
(.docx/.pdf/.xlsx; slower) — off by default.

    python "Database Codes/gen_mkt_doc_summaries.py"            # all categories
    python "Database Codes/gen_mkt_doc_summaries.py" NPRR PGRR  # subset
"""
import os
import re
import sys
import json
from urllib.parse import quote

ROOT = os.path.join("Documents Database", "ERCOT.MKT.RULES")
WEB_BASE = "/Power.Talks/Documents%20Database/ERCOT.MKT.RULES"

DOC_EXTS = (".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx")
_MONTHS = "jan feb mar apr may jun jul aug sep oct nov dec".split()

# Document-type detection from the filename (first match wins, order matters).
DOC_TYPES = [
    ("Impact Analysis",          r"impact[_ ]?analysis"),
    ("Cost/Benefit Analysis",    r"\bcba\b|cost[_ ]?benefit"),
    ("Board Action Report",      r"board[_ ]?action"),
    ("TAC Recommendation Report",r"tac[_ ]?rec"),
    ("PRS Recommendation Report",r"prs[_ ]?rec"),
    ("ROS Report",               r"ros[_ ]?(report|rec)"),
    ("WMS Report",               r"wms[_ ]?(report|rec)"),
    ("Recommendation Report",    r"recommendation|rec[_ ]?report"),
    ("Ballot",                   r"\bballot\b"),
    ("Comments",                 r"\bcomment"),
    ("Markup / Redline",         r"markup|redline|black[_ ]?line|black[_ ]?lined"),
    ("Presentation",             r"presentation|slides|overview|update"),
    ("Impact Analysis",          r"\bia\b"),
    ("As-Built",                 r"as[_ ]?built"),
]


def detect_doc_type(name, seq):
    low = name.lower().replace("_", " ")   # '_' is a regex word char — break it so \b works
    for label, pat in DOC_TYPES:
        if re.search(pat, low):
            return label
    if seq == 1:                       # the first item is usually the request itself
        return "Revision Request"
    return "Document"


def detect_date(name):
    digits = re.findall(r"\d+", name)
    for d in digits:
        if len(d) == 8:                # YYYYMMDD or MMDDYYYY
            if d[:4].startswith(("19", "20")):
                y, m, dd = d[:4], d[4:6], d[6:8]
            else:
                m, dd, y = d[:2], d[2:4], d[4:8]
            if 1 <= int(m) <= 12 and 1 <= int(dd) <= 31:
                return f"{y}-{m}-{dd}"
        if len(d) == 6:                # MMDDYY
            m, dd, yy = d[:2], d[2:4], d[4:6]
            if 1 <= int(m) <= 12 and 1 <= int(dd) <= 31:
                yr = ("20" if int(yy) < 50 else "19") + yy
                return f"{yr}-{m}-{dd}"
    return None


def detect_author(name):
    low = name.lower().replace("_", " ")
    for token, label in [("ercot", "ERCOT"), ("imm", "IMM"), ("opuc", "OPUC"),
                         ("puct", "PUCT"), ("tcpa", "TCPA"), ("tiec", "TIEC")]:
        if re.search(rf"\b{token}\b", low):
            return label
    return None


def clean_title(name, issue_id):
    t = os.path.splitext(name)[0]
    # drop the leading "<num><cat>_NN_" / "<cat><num>_NN_" packet prefix
    t = re.sub(r"^\d+[a-z]{2,7}[_ -]?\d{1,3}[_ -]+", "", t, flags=re.I)
    t = re.sub(rf"^{re.escape(issue_id)}[_ -]+", "", t, flags=re.I)
    t = re.sub(r"[_]+", " ", t)
    t = re.sub(r"\b\d{6,8}\b", "", t)               # date blobs
    t = re.sub(r"\b(rev\s*\d+|v\d+|final|clean|redline)\b", "", t, flags=re.I)
    t = re.sub(r"\s+", " ", t).strip(" -—_.")
    return t.title() if t and t.islower() else t


def summarize_file(fname, issue_id, seq):
    doc_type = detect_doc_type(fname, seq)
    date = detect_date(fname)
    author = detect_author(fname)
    title = clean_title(fname, issue_id)
    bits = [f"{doc_type} for {issue_id}"]
    if author:
        bits[0] = f"{author} {doc_type.lower()} for {issue_id}"
    if title and title.lower() not in doc_type.lower():
        bits.append(f"— {title}")
    if date:
        bits.append(f"(dated {date})")
    summary = " ".join(bits).strip() + "."
    # Title Name = the filename interpreted as "[Title Name].[file format]" — the
    # stem before the extension, with separators turned into spaces for display.
    stem = os.path.splitext(fname)[0]
    title_name = re.sub(r"\s+", " ", re.sub(r"[_-]+", " ", stem)).strip()
    return {
        "file": fname,
        "title": title_name,
        "doc_type": doc_type,
        "date": date,
        "date_posted": date,        # the document's own date (right-panel card)
        "author": author,
        "submitter": author,        # who filed it; AI mode may refine
        "summary": summary,
        "key_points": [],
        # Per-document report fields (filled by --ai mode reading the document).
        "revision_reason": None,
        "description": None,
        "justification": None,
        "detailed_background": None,
        "download_url": None,        # set by build_source_documents (needs cat/issue)
    }


def seq_key(fname):
    """Sort key: the sequence number that follows '[rule#][CAT]-' / '_NN_',
    e.g. '1264NPRR-19 …' -> 19, '100nprr_01_…' -> 1. Unmatched files sort last."""
    m = re.match(r"^\d+[a-z]+[-_ ]+(\d+)", fname, re.I)
    return (int(m.group(1)) if m else 9999, fname.lower())


# ── AI per-document report (--ai): read the document, ask Claude ─────────────
AI_MODEL = "claude-haiku-4-5-20251001"
AI_MAX_TOKENS = 1200
_REPORT_KEYS = ("revision_reason", "description", "justification",
                "detailed_background", "submitter")
_SUMM_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "summarize_MKT_Rules")
_ai_client = None
_word_app = None
_ai_usage = {"calls": 0, "in": 0, "out": 0}


def _get_ai():
    global _ai_client
    if _ai_client is None:
        import anthropic
        if _SUMM_DIR not in sys.path:
            sys.path.insert(0, _SUMM_DIR)
        from anthropic_key import get_anthropic_key
        _ai_client = anthropic.Anthropic(api_key=get_anthropic_key())
    return _ai_client


def _quit_word():
    global _word_app
    if _word_app is not None:
        try:
            _word_app.Quit()
        except Exception:
            pass
        _word_app = None


def extract_text_any(path):
    """Best-effort plain text from a document (.docx/.doc/.pdf/.xls/.xlsx)."""
    global _word_app
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            from docx import Document
            d = Document(path)
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for tbl in d.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
            return "\n".join(parts)
        if ext == ".doc":
            import win32com.client
            if _word_app is None:
                _word_app = win32com.client.Dispatch("Word.Application")
                _word_app.Visible = False
            doc = _word_app.Documents.Open(os.path.abspath(path), ReadOnly=True)
            t = doc.Content.Text
            doc.Close(False)
            return t
        if ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return "\n".join((pg.extract_text() or "") for pg in pdf.pages)
        if ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True, data_only=True)
            rows = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        rows.append(" | ".join(cells))
            return "\n".join(rows)
        if ext == ".xls":
            import xlrd
            bk = xlrd.open_workbook(path)
            rows = []
            for sh in bk.sheets():
                for r in range(sh.nrows):
                    cells = [str(sh.cell_value(r, c)).strip() for c in range(sh.ncols)
                             if str(sh.cell_value(r, c)).strip()]
                    if cells:
                        rows.append(" | ".join(cells))
            return "\n".join(rows)
    except Exception as e:
        print(f"    [READ-ERR] {os.path.basename(path)}: {e}")
    return ""


def ai_document_report(text, issue_id, fname, doc_type):
    """Ask Claude to summarize ONE document into the per-document report fields.
    Returns a dict over _REPORT_KEYS (values or None); {} if nothing usable."""
    text = (text or "").strip()
    if len(text) < 40:
        return {}
    prompt = (
        f"You are reading ONE document submitted in the ERCOT stakeholder process "
        f"for revision request {issue_id}.\n"
        f"Document file: {fname}\nDocument type: {doc_type}\n\n"
        "ERCOT revision-request forms label their sections explicitly — look for "
        "headings/table rows such as 'Reason for Revision', 'Revision Description', "
        "'Justification', 'Business Case', 'Impact Analysis'. Extract from those "
        "wherever present.\n\n"
        "Return a JSON object with these keys, drawn from THIS document's text:\n"
        '- "revision_reason": why the revision is needed (the Reason for Revision)\n'
        '- "description": what the document revises/proposes (the Revision Description)\n'
        '- "justification": the business case / justification for the change\n'
        '- "detailed_background": one or two plain-English paragraphs giving the '
        "background of the changes this document covers — what is changing and why\n"
        '- "submitter": person and/or company that authored/filed this document, if stated\n\n'
        "Fill every field the document supports; use null ONLY when the document "
        "genuinely has nothing for it. Do not invent facts. Plain text, no markdown. "
        "Return ONLY the JSON object.\n\n"
        f"DOCUMENT TEXT:\n{text[:14000]}"
    )
    try:
        msg = _get_ai().messages.create(
            model=AI_MODEL, max_tokens=AI_MAX_TOKENS,
            messages=[{"role": "user", "content": prompt}])
        _ai_usage["calls"] += 1
        _ai_usage["in"] += msg.usage.input_tokens
        _ai_usage["out"] += msg.usage.output_tokens
        raw = msg.content[0].text.strip()
        m = re.search(r"\{.*\}", raw, re.S)
        data = json.loads(m.group(0) if m else raw)
        return {k: (data.get(k) or None) for k in _REPORT_KEYS}
    except Exception as e:
        print(f"    [AI-ERR] {fname}: {e}")
        return {}


def _aslist(v):
    if isinstance(v, list):
        return "; ".join(str(x) for x in v if x)
    return v or None


def build_source_documents(cat, issue_id, folder, ai=False, profile=None):
    profile = profile or {}
    # The primary revision-request submission carries the issue's own
    # Reason/Description/Business Case — already extracted into the Profile by the
    # profile generator. Use them to fill the -01's report fields the AI leaves blank.
    prof_fallback = {
        "revision_reason": _aslist(profile.get("reason_for_revision")),
        "description": profile.get("revision_description"),
        "justification": profile.get("business_case"),
    }
    files = sorted(
        (f for f in os.listdir(folder)
         if os.path.isfile(os.path.join(folder, f))
         and f.lower().endswith(DOC_EXTS)
         and not f.lower().endswith((".zip", ".tmp"))),
        key=seq_key,
    )
    docs = []
    for i, f in enumerate(files, 1):
        entry = summarize_file(f, issue_id, i)
        entry["download_url"] = f"{WEB_BASE}/{quote(cat)}/{quote(issue_id)}/{quote(f)}"
        if ai:
            rep = ai_document_report(extract_text_any(os.path.join(folder, f)),
                                     issue_id, f, entry["doc_type"])
            for k in _REPORT_KEYS:
                if rep.get(k):
                    entry[k] = rep[k]
            if rep.get("submitter"):
                entry["submitter"] = rep["submitter"]
            print(f"    [AI] {f}")
        # Primary submission (-01 / Revision Request): backfill any blank
        # Reason/Description/Justification from the issue Profile.
        if i == 1 or entry["doc_type"] == "Revision Request":
            for k, v in prof_fallback.items():
                if not entry.get(k) and v:
                    entry[k] = v
        docs.append(entry)
    return docs


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    ai = "--ai" in sys.argv
    # Positional args that look like an issue id (e.g. NPRR1340) target single
    # issues; bare category names (e.g. NPRR) limit to a category.
    issue_targets = {a.upper() for a in args if re.match(r"^[A-Za-z]+\d+$", a)}
    cat_only = {a.upper() for a in args if not re.match(r"^[A-Za-z]+\d+$", a)}
    updated = skipped = 0
    try:
        for cat in sorted(os.listdir(ROOT)):
            if cat_only and cat.upper() not in cat_only:
                continue
            cdir = os.path.join(ROOT, cat)
            if not os.path.isdir(cdir):
                continue
            for issue_id in sorted(os.listdir(cdir)):
                if issue_targets and issue_id.upper() not in issue_targets:
                    continue
                folder = os.path.join(cdir, issue_id)
                if not os.path.isdir(folder):
                    continue
                profile = os.path.join(folder, "Quick runs", f"{issue_id} Profile.json")
                if not os.path.exists(profile):
                    skipped += 1
                    continue
                try:
                    with open(profile, encoding="utf-8") as fh:
                        data = json.load(fh)
                    if ai:
                        print(f"  {issue_id} …")
                    data["source_documents"] = build_source_documents(cat, issue_id, folder, ai=ai, profile=data)
                    with open(profile, "w", encoding="utf-8") as fh:
                        json.dump(data, fh, indent=2, ensure_ascii=False)
                    updated += 1
                    if not ai and updated % 200 == 0:
                        print(f"  ... {updated} profiles updated")
                except Exception as e:
                    print(f"  ERROR {cat}/{issue_id}: {e}")
    finally:
        _quit_word()
    print(f"\nDone. profiles updated={updated}  (no profile){skipped}")
    if _ai_usage["calls"]:
        print(f"AI: {_ai_usage['calls']} calls, in={_ai_usage['in']} out={_ai_usage['out']} tokens ({AI_MODEL})")


if __name__ == "__main__":
    main()
