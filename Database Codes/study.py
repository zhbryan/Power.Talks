"""
study.py — Power.Talks
Reads newly downloaded documents from Documents Database/, sends each to the
Claude API for summarization using the Default Studies prompt, and saves
Word (.docx) reports to Reports Database/.
"""

import os
import json
import logging
import re
from datetime import datetime
from pathlib import Path

import anthropic
import pdfplumber
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(BASE_DIR, "Documents Database")
REPORTS_DIR = os.path.join(BASE_DIR, "Reports Database")
INDEX_FILE = os.path.join(BASE_DIR, "Database Codes", "processed_files.json")
MANIFEST_FILE = os.path.join(BASE_DIR, "html", "reports-manifest.json")

MODEL = "claude-sonnet-4-6"
MAX_CHARS = 150_000  # Truncate very large documents to stay within context limits

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

DEFAULT_STUDIES_PROMPT = """You are analyzing an ERCOT regulatory document for the Power.Talks project.

Produce a structured report with exactly these five sections:

### 1. Document Title & Date
State the full document title and its publication or effective date.

### 2. Executive Summary
Write 3–5 sentences capturing the document's core purpose and most important outcome.

### 3. Key Points
Bullet list of the most significant facts, rule changes, or decisions in the document.

### 4. Market Impact
Describe who is affected (Market Participants, QSEs, TSPs, REPs, etc.) and how. Be specific about operational or financial implications.

### 5. Action Items & Deadlines
List any compliance requirements, filing deadlines, or implementation milestones. If none, state "No action items identified."

Formatting rules:
- Use clear, plain business English
- Keep bullet points concise (one idea per bullet)
- Dates must be in MM/DD/YYYY format

Here is the document text:
"""

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(BASE_DIR, "Database Codes", "study.log")),
        logging.StreamHandler(),
    ],
)
log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_index() -> dict:
    if os.path.exists(INDEX_FILE):
        with open(INDEX_FILE, "r") as f:
            return json.load(f)
    return {}


def save_index(index: dict):
    with open(INDEX_FILE, "w") as f:
        json.dump(index, f, indent=2)


def extract_text_pdf(path: str) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_docx(path: str) -> str:
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(path)
    elif ext in {".docx", ".doc"}:
        return extract_text_docx(path)
    return ""


def summarize(text: str, filename: str) -> str:
    client = anthropic.Anthropic()
    prompt = DEFAULT_STUDIES_PROMPT + text[:MAX_CHARS]
    message = client.messages.create(
        model=MODEL,
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text


def derive_report_name(rel_path: str, source_filename: str) -> str:
    parts = Path(rel_path).parts
    # parts[0] = "ERCOT.MKT.RULES" or "ERCOT.STKHDR.MEETS"
    # parts[1] = category like "NPRR", "BOARD", etc.
    category = parts[1] if len(parts) >= 2 else "UNKNOWN"
    stem = Path(source_filename).stem
    date_str = datetime.now().strftime("%Y%m%d")
    return f"{category}_{stem}_{date_str}.docx"


def write_report(summary_text: str, report_path: str, source_filename: str):
    doc = DocxDocument()

    # Title
    title = doc.add_heading("Power.Talks — ERCOT Document Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Source: {source_filename}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Parse sections from Claude's markdown output
    section_pattern = re.compile(r"###\s+\d+\.\s+(.+)")
    lines = summary_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        match = section_pattern.match(line.strip())
        if match:
            doc.add_heading(match.group(1).strip(), level=2)
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            para = doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())
        i += 1

    doc.save(report_path)


def update_manifest(manifest: list, entry: dict):
    manifest.append(entry)
    os.makedirs(os.path.dirname(MANIFEST_FILE), exist_ok=True)
    with open(MANIFEST_FILE, "w") as f:
        json.dump(manifest, f, indent=2)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def run():
    os.makedirs(REPORTS_DIR, exist_ok=True)
    os.makedirs(os.path.join(BASE_DIR, "html"), exist_ok=True)

    index = load_index()
    manifest = []
    if os.path.exists(MANIFEST_FILE):
        with open(MANIFEST_FILE, "r") as f:
            manifest = json.load(f)

    processed_count = 0
    log.info("=== Study started: %s ===", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    for root, _, files in os.walk(DOCS_DIR):
        for filename in files:
            ext = Path(filename).suffix.lower()
            if ext not in SUPPORTED_EXTENSIONS:
                continue

            full_path = os.path.join(root, filename)
            rel_path = os.path.relpath(full_path, DOCS_DIR)

            if rel_path in index:
                log.debug("Already processed, skipping: %s", rel_path)
                continue

            log.info("Processing: %s", rel_path)

            text = extract_text(full_path)
            if not text.strip():
                log.warning("No text extracted from %s, skipping.", filename)
                index[rel_path] = {"status": "no_text", "processed_at": datetime.now().isoformat()}
                save_index(index)
                continue

            try:
                summary = summarize(text, filename)
            except Exception as e:
                log.error("Claude API error for %s: %s", filename, e)
                index[rel_path] = {"status": "error", "error": str(e), "processed_at": datetime.now().isoformat()}
                save_index(index)
                continue

            report_name = derive_report_name(rel_path, filename)
            report_path = os.path.join(REPORTS_DIR, report_name)
            write_report(summary, report_path, filename)

            index[rel_path] = {"status": "done", "report": report_name, "processed_at": datetime.now().isoformat()}
            save_index(index)

            # Extract category for manifest
            parts = Path(rel_path).parts
            category = parts[1] if len(parts) >= 2 else "UNKNOWN"
            parent = parts[0] if len(parts) >= 1 else ""
            group = "MKT.RULES" if "MKT.RULES" in parent else "STKHDR.MEETS"

            manifest_entry = {
                "report_file": report_name,
                "source_file": filename,
                "category": category,
                "group": group,
                "date": datetime.now().strftime("%Y-%m-%d"),
                "summary_snippet": summary[:200].replace("\n", " ") + "...",
            }
            update_manifest(manifest, manifest_entry)
            manifest.append(manifest_entry)  # keep in-memory list in sync

            processed_count += 1
            log.info("Report saved: %s", report_name)

    log.info("=== Study finished. Reports generated: %d ===", processed_count)
    return processed_count


if __name__ == "__main__":
    run()
