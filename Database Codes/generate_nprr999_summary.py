"""Generate NPRR999 Summary.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

TERRACOTTA = RGBColor(0xC8, 0x62, 0x3E)
DARK       = RGBColor(0x1B, 0x1A, 0x17)
MUTED      = RGBColor(0x7A, 0x74, 0x6A)


def h1(doc, text):
    p = doc.add_heading("", level=1)
    p.clear()
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(20)
    run.font.color.rgb = TERRACOTTA
    run.bold = False
    return p


def h2(doc, text):
    p = doc.add_heading("", level=2)
    p.clear()
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    run.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def meta_row(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10)
    r1.font.color.rgb = MUTED
    r2 = p.add_run(value or "—")
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4CFC9")
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = Inches(w)
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    return tbl


# ── Cover ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("ERCOT  ·  Nodal Protocol Revision Request")
r.font.name = "Calibri"
r.font.size = Pt(9)
r.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(2)

h1(doc, "NPRR999 — DC Tie Ramp Limitations")

p = doc.add_paragraph()
r = p.add_run("Summary Report")
r.font.name = "Calibri"
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(4)

rule(doc)

# ── 1. Executive Summary ──────────────────────────────────────────────────────
h2(doc, "1. Executive Summary")
body(doc,
    "NPRR999 proposes targeted revisions to Protocol Section 4.4.4 (DC Tie Schedules) and "
    "introduces new Section 4.4.4.3 (Management of DC Tie Schedules due to Ramp Limitations). "
    "The revision codifies ERCOT's authority and procedure when system conditions show "
    "insufficient ramp capability to meet the aggregate scheduled ramp of all Direct Current "
    "(DC) Ties: ERCOT shall curtail DC Tie Schedules on a last-in, first-out (LIFO) basis, "
    "but may first request Qualified Scheduling Entities (QSEs) to voluntarily resubmit "
    "e-Tags with an adjusted ramp duration if sufficient time allows.")

h2(doc, "1.1  Background")
body(doc,
    "The NPRR was directly triggered by a May 23, 2017 PUCT order in Project No. 46304 "
    "(Oversight Proceeding arising from PUC Docket No. 45624 — the Southern Cross "
    "Transmission DC Tie). Directive 3 of that order required ERCOT to determine what ramp "
    "rate restrictions would be necessary to accommodate interconnection of the proposed "
    "Southern Cross DC Tie.")
body(doc,
    "The Southern Cross DC Tie could swing up to 4,100 MW between Operating Hours — far "
    "exceeding the swings from all existing ERCOT DC Ties, which typically ramp over ten "
    "minutes. While existing ten-minute ramps have historically been sufficient, ERCOT "
    "determined this would be inadequate for Southern Cross. Independently, NERC Reliability "
    "Standard INT-006-4 R1 already requires ERCOT to reject or curtail any DC Tie Schedule "
    "it does not expect to be supportable in magnitude or ramp; however, the Protocols did "
    "not clearly describe how ERCOT would operationally address an identified DC Tie ramp "
    "limitation. NPRR999 fills that gap.")

h2(doc, "1.2  Key Protocol Change")
body(doc,
    "Following ERCOT comments filed July 2, 2020, the final approved language clarifies that "
    "ERCOT's ramp assessment focuses on preserving sufficient Physical Responsive Capability "
    "(PRC) to avoid entering an Energy Emergency Alert (EEA) Level 1. When a DC Tie ramp "
    "limitation is identified, ERCOT may request voluntary e-Tag resubmission (if time "
    "permits); curtailment is applied on a system-wide LIFO basis — not tied to a "
    "specific DC Tie — and QSEs are under no obligation to resubmit.")

h2(doc, "1.3  Potential Impacts")
impacts = [
    ("ERCOT Operations",
     "New desk procedures and operator training required. ERCOT retains discretion over "
     "which QSE to contact for voluntary resubmission requests."),
    ("QSEs / DC Tie Operators",
     "Exposed to curtailment risk under a system-wide LIFO basis when DC Tie ramp "
     "limitations are triggered. Voluntary resubmission requests impose no obligation "
     "but may be received on short notice."),
    ("Energy Management System (EMS)",
     "System changes required — 100% ERCOT effort. Implementation is contingent on "
     "the Intra-Hour Variability (iCAT) Tool project."),
    ("Credit & Settlement",
     "ERCOT Credit Staff and the Credit Work Group confirmed no changes to credit "
     "monitoring activity or liability calculation are required."),
    ("Cost",
     "Estimated at less than $5,000, absorbed within O&M budgets of affected departments."),
]
for label, text in impacts:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)

rule(doc)

# ── 2. NPRR Details ───────────────────────────────────────────────────────────
h2(doc, "2. NPRR Details")
meta_row(doc, "NPRR Number",          "999")
meta_row(doc, "Title",                "DC Tie Ramp Limitations")
meta_row(doc, "Status",               "Approved")
meta_row(doc, "Date Posted",          "February 20, 2020")
meta_row(doc, "Effective Date",       "Upon project implementation of the Intra-Hour Variability (iCAT) Tool")
meta_row(doc, "Market Segment",       "Not Applicable")
meta_row(doc, "Protocol Sections",    "4.4.4 – DC Tie Schedules (revised); 4.4.4.3 – Management of DC Tie Schedules due to Ramp Limitations (new)")
meta_row(doc, "Related Revisions",    "NPRR1008 (RTC – NP 4: Day-Ahead Operations) also revises Section 4.4.4 — coordinate implementation")
meta_row(doc, "Sponsor",              "Nitika Mago  ·  ERCOT  ·  nitika.mago@ercot.com  ·  512-248-6601")
meta_row(doc, "Market Rules Contact", "Brittney Albracht  ·  ERCOT  ·  Brittney.Albracht@ercot.com  ·  512-225-7027")

rule(doc)

# ── 3. Impact Analysis ────────────────────────────────────────────────────────
h2(doc, "3. Impact Analysis")

body(doc, "Initial Impact Analysis — February 20, 2020")
make_table(doc,
    ["Category", "Detail"],
    [
        ("Estimated Cost",          "Less than $5,000; absorbed within O&M budgets of affected departments."),
        ("Time Requirements",       "No project required. Can take effect upon ERCOT Board approval."),
        ("ERCOT Staffing",          "No staffing impacts."),
        ("ERCOT Computer Systems",  "No system impacts."),
        ("Business Function",       "No business function impacts."),
        ("Grid Operations",         "ERCOT will update grid operations and practices to implement this NPRR."),
        ("Comments",                "Desk procedure update and operator training required."),
    ],
    [2.0, 4.5]
)

doc.add_paragraph()

body(doc, "Revised Impact Analysis — August 28, 2020")
make_table(doc,
    ["Category", "Detail"],
    [
        ("Estimated Cost",          "Less than $5,000; absorbed within O&M budgets."),
        ("Time Requirements",       "Dependent on iCAT Tool project implementation (no standalone project required)."),
        ("ERCOT Staffing",          "Implementation Labor: 100% ERCOT / 0% Vendor. No ongoing staffing impacts."),
        ("ERCOT Computer Systems",  "Energy Management System (EMS) — 100% impacted."),
        ("Business Function",       "No business function impacts."),
        ("Grid Operations",         "ERCOT will update grid operations and practices to implement this NPRR."),
        ("Comments",                "NPRR999 implementation will be dependent on the internal ERCOT iCAT project."),
    ],
    [2.0, 4.5]
)

rule(doc)

# ── 4. Stakeholder Discussion Timeline ───────────────────────────────────────
h2(doc, "4. Stakeholder Discussion Timeline")

timeline_rows = [
    ("2020-02-20", "Sponsor / ERCOT", "NPRR Posted",
     "NPRR999 posted per PUCT Project No. 46304 Directive 3. Southern Cross DC Tie could cause 4,100 MW swings between Operating Hours."),
    ("2020-02-20", "ERCOT Staff",     "Impact Analysis Filed",
     "Initial IA published. Estimated cost <$5k; no project or system changes required."),
    ("2020-06-11", "PRS",             "Motion to Table — Carries (Unanimous)",
     "Participants discussed tabling pending ERCOT staff comments on the revision language."),
    ("2020-07-02", "ERCOT Staff",     "ERCOT Comments Filed",
     "Revised language: (1) LIFO curtailment made primary; (2) voluntary e-Tag resubmission added as prior step; (3) ramp assessment tied to PRC / EEA Level 1 threshold; (4) NPRR1008 conflict flagged."),
    ("2020-07-16", "PRS",             "Recommend Approval (as amended 7/2) — Carries (Unanimous)",
     "Participants reviewed and accepted 7/2/20 ERCOT comments. All Market Segments present."),
    ("2020-08-13", "PRS",             "Motion to Table — Carries (1 IOU abstention: Lone Star Transmission)",
     "ERCOT Staff requested additional time to review Impact Analysis due to new EMS dependency."),
    ("2020-08-28", "ERCOT Staff",     "Revised Impact Analysis Filed",
     "Updated to reflect EMS dependency (100% ERCOT). Effective date now conditional on iCAT Tool implementation."),
    ("2020-09-10", "PRS",             "Endorse & Forward to TAC (8/13 PRS Report + Revised IA) — Carries (Unanimous)",
     "No substantive discussion. PRS endorsed revised IA and forwarded package to TAC."),
    ("2020-09-23", "TAC",             "Recommend Approval (per 9/10 PRS Report) — Carries (1 IPM abstention: Shell)",
     "No discussion. ERCOT formally stated support. Credit WG confirmed no credit impacts."),
    ("2020-10-13", "ERCOT Board",     "Approved",
     "Board approved NPRR999 as recommended by TAC in the 9/23/20 TAC Report."),
]

make_table(doc,
    ["Date", "Body", "Action / Vote", "Notes"],
    timeline_rows,
    [1.0, 1.2, 2.1, 2.2]
)

rule(doc)

# ── 5. Current Status ─────────────────────────────────────────────────────────
h2(doc, "5. Current Status")
body(doc,
    "NPRR999 was approved by the ERCOT Board on October 13, 2020, as recommended by TAC. "
    "The effective date is contingent on project implementation of the Intra-Hour Variability "
    "(iCAT) Tool — an internal ERCOT system project. As of Board approval, the iCAT "
    "implementation timeline was designated as 'to be determined.'")
body(doc,
    "Credit Work Group review confirmed no changes to credit monitoring activity or liability "
    "calculation are required. NPRR1008 (RTC – NP 4: Day-Ahead Operations) was identified "
    "as a related revision also proposing changes to Section 4.4.4; implementation coordination "
    "with NPRR1008 should be verified before the iCAT project goes live.")

# ── Save .docx ────────────────────────────────────────────────────────────────
out_dir  = r"E:\wamp64\www\Power.Talks\Documents Database\ERCOT.MKT.RULES\NPRR\NPRR999\Quick runs"
out_path = os.path.join(out_dir, "NPRR999 Summary.docx")
os.makedirs(out_dir, exist_ok=True)
doc.save(out_path)
print(f"Saved: {out_path}")

# ── Save companion Summary.json (consumed by the web UI) ─────────────────────
import json

summary_data = {
    "nprr_number": 999,
    "title": "DC Tie Ramp Limitations",
    "status": "Approved",
    "date_posted": "2020-02-20",
    "effective_date": "Upon project implementation of the Intra-Hour Variability (iCAT) Tool",
    "protocol_sections": [
        "4.4.4 – DC Tie Schedules (revised)",
        "4.4.4.3 – Management of DC Tie Schedules due to Ramp Limitations (new)",
    ],
    "sponsor": "Nitika Mago · ERCOT · nitika.mago@ercot.com · 512-248-6601",
    "executive_summary": (
        "NPRR999 proposes targeted revisions to Protocol Section 4.4.4 (DC Tie Schedules) and "
        "introduces new Section 4.4.4.3 (Management of DC Tie Schedules due to Ramp Limitations). "
        "The revision codifies ERCOT's authority and procedure when system conditions show "
        "insufficient ramp capability to meet the aggregate scheduled ramp of all DC Ties: "
        "ERCOT shall curtail DC Tie Schedules on a last-in, first-out (LIFO) basis, but may "
        "first request QSEs to voluntarily resubmit e-Tags with an adjusted ramp duration if "
        "sufficient time allows."
    ),
    "background": (
        "The NPRR was directly triggered by a May 23, 2017 PUCT order in Project No. 46304 "
        "(Oversight Proceeding arising from PUC Docket No. 45624 — the Southern Cross "
        "Transmission DC Tie). Directive 3 required ERCOT to determine ramp rate restrictions "
        "for the proposed Southern Cross DC Tie, which could swing up to 4,100 MW between "
        "Operating Hours — far exceeding existing ERCOT DC Ties that ramp over ten minutes. "
        "NERC Reliability Standard INT-006-4 R1 already required ERCOT to reject/curtail "
        "unsupportable DC Tie Schedules, but the Protocols lacked clear operational language. "
        "NPRR999 fills that gap."
    ),
    "key_change": (
        "Following ERCOT comments filed July 2, 2020, the final approved language clarifies "
        "that ERCOT's ramp assessment focuses on preserving sufficient Physical Responsive "
        "Capability (PRC) to avoid entering an Energy Emergency Alert (EEA) Level 1. When a "
        "DC Tie ramp limitation is identified, ERCOT may request voluntary e-Tag resubmission "
        "(if time permits); curtailment is applied system-wide on a LIFO basis — not tied to "
        "a specific DC Tie — and QSEs are under no obligation to resubmit."
    ),
    "impacts": [
        {"category": "ERCOT Operations",
         "text": "New desk procedures and operator training required. ERCOT retains discretion over which QSE to contact for voluntary resubmission requests."},
        {"category": "QSEs / DC Tie Operators",
         "text": "Exposed to curtailment risk under a system-wide LIFO basis when DC Tie ramp limitations are triggered. Voluntary resubmission requests impose no obligation but may be received on short notice."},
        {"category": "Energy Management System (EMS)",
         "text": "System changes required — 100% ERCOT effort. Implementation is contingent on the Intra-Hour Variability (iCAT) Tool project."},
        {"category": "Credit & Settlement",
         "text": "ERCOT Credit Staff and the Credit Work Group confirmed no changes to credit monitoring activity or liability calculation are required."},
        {"category": "Cost",
         "text": "Estimated at less than $5,000, absorbed within O&M budgets of affected departments."},
    ],
    "impact_analysis": [
        {
            "label": "Initial Impact Analysis — February 20, 2020",
            "rows": [
                ["Estimated Cost",         "Less than $5,000; absorbed within O&M budgets."],
                ["Time Requirements",      "No project required. Can take effect upon ERCOT Board approval."],
                ["ERCOT Staffing",         "No staffing impacts."],
                ["ERCOT Computer Systems", "No system impacts."],
                ["Business Function",      "No business function impacts."],
                ["Grid Operations",        "ERCOT will update grid operations and practices."],
                ["Comments",               "Desk procedure update and operator training required."],
            ],
        },
        {
            "label": "Revised Impact Analysis — August 28, 2020",
            "rows": [
                ["Estimated Cost",         "Less than $5,000; absorbed within O&M budgets."],
                ["Time Requirements",      "Dependent on iCAT Tool project implementation."],
                ["ERCOT Staffing",         "Implementation Labor: 100% ERCOT / 0% Vendor. No ongoing impacts."],
                ["ERCOT Computer Systems", "Energy Management System (EMS) — 100% impacted."],
                ["Business Function",      "No business function impacts."],
                ["Grid Operations",        "ERCOT will update grid operations and practices."],
                ["Comments",               "NPRR999 implementation will be dependent on the internal ERCOT iCAT project."],
            ],
        },
    ],
    "timeline": [
        {"date": "2020-02-20", "body": "Sponsor / ERCOT", "action": "NPRR Posted",
         "notes": "NPRR999 posted per PUCT Project No. 46304 Directive 3. Southern Cross DC Tie could cause 4,100 MW swings between Operating Hours."},
        {"date": "2020-02-20", "body": "ERCOT Staff",     "action": "Impact Analysis Filed",
         "notes": "Initial IA published. Estimated cost <$5k; no project or system changes required."},
        {"date": "2020-06-11", "body": "PRS",             "action": "Motion to Table — Carries (Unanimous)",
         "notes": "Participants discussed tabling pending ERCOT staff comments on the revision language."},
        {"date": "2020-07-02", "body": "ERCOT Staff",     "action": "ERCOT Comments Filed",
         "notes": "Revised language: LIFO curtailment made primary; voluntary e-Tag resubmission added as prior step; ramp assessment tied to PRC / EEA Level 1; NPRR1008 conflict flagged."},
        {"date": "2020-07-16", "body": "PRS",             "action": "Recommend Approval (as amended 7/2) — Carries (Unanimous)",
         "notes": "Participants reviewed and accepted 7/2/20 ERCOT comments. All Market Segments present."},
        {"date": "2020-08-13", "body": "PRS",             "action": "Motion to Table — Carries (1 IOU abstention: Lone Star Transmission)",
         "notes": "ERCOT Staff requested additional time to review Impact Analysis due to new EMS dependency."},
        {"date": "2020-08-28", "body": "ERCOT Staff",     "action": "Revised Impact Analysis Filed",
         "notes": "Updated to reflect EMS dependency (100% ERCOT). Effective date now conditional on iCAT Tool implementation."},
        {"date": "2020-09-10", "body": "PRS",             "action": "Endorse & Forward to TAC — Carries (Unanimous)",
         "notes": "No substantive discussion. PRS endorsed revised IA and forwarded package to TAC."},
        {"date": "2020-09-23", "body": "TAC",             "action": "Recommend Approval — Carries (1 IPM abstention: Shell)",
         "notes": "No discussion. ERCOT formally stated support. Credit WG confirmed no credit impacts."},
        {"date": "2020-10-13", "body": "ERCOT Board",     "action": "Approved",
         "notes": "Board approved NPRR999 as recommended by TAC in the 9/23/20 TAC Report."},
    ],
    "current_status": [
        "NPRR999 was approved by the ERCOT Board on October 13, 2020, as recommended by TAC. "
        "The effective date is contingent on project implementation of the Intra-Hour Variability "
        "(iCAT) Tool — an internal ERCOT system project. As of Board approval, the iCAT "
        "implementation timeline was designated as 'to be determined.'",
        "Credit Work Group review confirmed no changes to credit monitoring activity or liability "
        "calculation are required. NPRR1008 (RTC – NP 4: Day-Ahead Operations) was identified "
        "as a related revision also proposing changes to Section 4.4.4; implementation coordination "
        "with NPRR1008 should be verified before the iCAT project goes live.",
    ],
}

json_path = os.path.join(out_dir, "NPRR999 Summary.json")
with open(json_path, "w", encoding="utf-8") as f:
    json.dump(summary_data, f, indent=2, ensure_ascii=False)
print(f"Saved: {json_path}")
