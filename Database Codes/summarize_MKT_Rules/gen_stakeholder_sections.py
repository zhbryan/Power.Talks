#!/usr/bin/env python3
"""
Generate the two comment-derived homepage sections for market-rule issues:

  * stakeholder_key_debates  — key positions / points of contention across the
                               comment documents filed by market participants.
  * ercot_imm_opinions       — ERCOT staff and Independent Market Monitor (IMM /
                               Potomac Economics) opinions.

Both are written into each issue's Summary.json. Content is AI-summarized from
the actual submitted comment documents; when an issue has no relevant comments
the field is left "" (the homepage then shows a blank section, by design).

Usage:
  python gen_stakeholder_sections.py NPRR 1238 1200 1345      # specific issues
  python gen_stakeholder_sections.py NPRR --recent 8          # newest 8 with comments
  python gen_stakeholder_sections.py NPRR --recent 8 --dry    # no writes, print only

Reads .docx (python-docx) and legacy .doc (Word COM). Requires the anthropic key.
"""

import os
import re
import sys
import json

from docx import Document
import win32com.client
import anthropic

try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
except Exception:
    pass

ROOT = r"E:\wamp64\www\Power.Talks\Documents Database\ERCOT.MKT.RULES"
AI_MODEL = "claude-haiku-4-5-20251001"
AI_MAX_TOKENS = 700
DOC_CHAR_CAP = 4000      # per-document text fed to the model
MAX_DOCS = 12            # cap documents per section

_ai = {"input": 0, "output": 0, "calls": 0}
_client = None
_word = None


def get_ai():
    global _client
    if _client is None:
        from anthropic_key import get_anthropic_key
        _client = anthropic.Anthropic(api_key=get_anthropic_key())
    return _client


def get_word():
    global _word
    if _word is None:
        _word = win32com.client.Dispatch("Word.Application")
        _word.Visible = False
    return _word


def quit_word():
    global _word
    if _word is not None:
        try:
            _word.Quit()
        except Exception:
            pass
        _word = None


def read_doc_text(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == '.docx':
            d = Document(path)
            parts = [p.text for p in d.paragraphs]
            for tbl in d.tables:
                for row in tbl.rows:
                    parts.append(' | '.join(c.text for c in row.cells))
            return '\n'.join(x for x in parts if x and x.strip())
        if ext == '.doc':
            wd = get_word()
            doc = wd.Documents.Open(os.path.abspath(path), ReadOnly=True)
            try:
                return doc.Content.Text.replace('\r', '\n')
            finally:
                doc.Close(False)
    except Exception as e:
        print(f"    (could not read {os.path.basename(path)}: {e})")
    return ""


# Party name = text between the "<num><CAT>-<seq> " prefix and " Comments".
_PARTY_PAT = re.compile(r'^\s*\d+[A-Za-z]+[-_ ]+\d+\s+(.*?)\s+comments\b', re.IGNORECASE)
_IMM_PAT = re.compile(r'\b(imm|independent market monitor|potomac)\b', re.IGNORECASE)


def party_of(fname):
    m = _PARTY_PAT.search(fname or "")
    return (m.group(1).strip() if m else "").strip()


def classify(fname):
    """Return 'ercot_imm', 'stakeholder', or None (not a comment doc)."""
    f = (fname or "").lower()
    if 'comment' not in f:
        return None
    party = party_of(fname)
    if not party:
        return None
    if _IMM_PAT.search(party):
        return 'ercot_imm'
    if party.strip().lower() == 'ercot':
        return 'ercot_imm'
    return 'stakeholder'


def _clean_ai(text):
    """Strip markdown the model sometimes emits despite instructions — the
    homepage renders plain text, so headings/bold must not leak through."""
    if not text:
        return ""
    lines = [ln for ln in text.splitlines() if not ln.lstrip().startswith('#')]
    t = '\n'.join(lines)
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)          # unbold
    t = re.sub(r'^\s*[-*]\s+', '', t, flags=re.M)   # de-bullet
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()


def ai_summarize(kind, issue_id, title, blocks):
    if not blocks:
        return ""
    joined = "\n\n".join(f"=== {p} ({d}) ===\n{t[:DOC_CHAR_CAP]}"
                         for p, d, t in blocks[:MAX_DOCS])
    if kind == 'stakeholder':
        ask = ("Summarize the KEY DEBATES and differing positions among market "
               "participants in 3-6 plain-English sentences — who supported or "
               "opposed what, the main points of contention, and any compromises. "
               "If the comments are purely procedural with no substantive debate, "
               "say so in one sentence.")
    else:
        ask = ("Summarize the opinions, concerns, and recommendations of ERCOT "
               "staff and/or the Independent Market Monitor (IMM) in 2-5 "
               "plain-English sentences. Distinguish ERCOT's position from the "
               "IMM's when both are present.")
    prompt = (f"The following are comment documents filed on ERCOT {issue_id} "
              f"({title or 'N/A'}).\n\n{joined}\n\n{ask} Write one or two plain-text "
              f"paragraphs only — no markdown, no headings, no bullet points, no "
              f"bold. Do not repeat the issue number.")
    try:
        msg = get_ai().messages.create(
            model=AI_MODEL, max_tokens=AI_MAX_TOKENS,
            messages=[{"role": "user", "content": prompt}])
        _ai["input"] += msg.usage.input_tokens
        _ai["output"] += msg.usage.output_tokens
        _ai["calls"] += 1
        return _clean_ai(msg.content[0].text)
    except Exception as e:
        print(f"    (AI failed: {e})")
        return ""


# ── ERCOT / IMM opinions: extracted from the "Opinion" table in Report docs ──
# PRS/TAC/Board/PUCT Reports carry rows "ERCOT Opinion" and "Independent Market
# Monitor Opinion". Early reports say "To be determined"; the latest report with
# a real value is authoritative, so scan in ascending sequence and let later
# values overwrite. Deterministic — no AI.
_OPIN_SKIP = re.compile(r'to be determined|^tbd$|^n/?a$|^none$|^\s*$', re.IGNORECASE)


def _seq(fname):
    m = re.match(r'^\s*\d+[A-Za-z]+[-_ ]+(\d+)', fname or "")
    return int(m.group(1)) if m else 9999


def doc_table_rows(path):
    ext = os.path.splitext(path)[1].lower()
    rows = []
    try:
        if ext == '.docx':
            d = Document(path)
            for tbl in d.tables:
                for row in tbl.rows:
                    rows.append([re.sub(r'\s+', ' ', c.text).strip() for c in row.cells])
        elif ext == '.doc':
            wd = get_word()
            doc = wd.Documents.Open(os.path.abspath(path), ReadOnly=True)
            try:
                for ti in range(1, doc.Tables.Count + 1):
                    t = doc.Tables(ti)
                    for r in range(1, t.Rows.Count + 1):
                        cells = []
                        for c in range(1, t.Columns.Count + 1):
                            try:
                                cells.append(t.Cell(r, c).Range.Text
                                             .replace('\r', ' ').replace('\x07', '').strip())
                            except Exception:
                                cells.append('')
                        rows.append(cells)
            finally:
                doc.Close(False)
    except Exception as e:
        print(f"    (could not read tables of {os.path.basename(path)}: {e})")
    return rows


def extract_opinions(base, sd):
    """Return (ercot_opinion, imm_opinion) from the latest Report doc that has
    real (non-'To be determined') values."""
    ercot = imm = ""
    docs = sorted(sd, key=lambda d: _seq((d.get('file') if isinstance(d, dict) else d) or ""))
    for doc in docs:
        fname = (doc.get('file') if isinstance(doc, dict) else doc) or ""
        if 'report' not in fname.lower() or not fname.lower().endswith(('.docx', '.doc')):
            continue
        for cells in doc_table_rows(os.path.join(base, fname)):
            if not cells:
                continue
            label = cells[0].lower()
            value = next((c for c in cells[1:] if c and c.lower() != label), "")
            if _OPIN_SKIP.search(value):
                continue
            if 'ercot opinion' in label:
                ercot = value
            elif 'market monitor opinion' in label:
                imm = value
    return ercot.strip(), imm.strip()


def process_issue(cat, folder, dry=False, force=False):
    base = os.path.join(ROOT, cat, folder)
    quick = os.path.join(base, "Quick runs")
    prof_path = os.path.join(quick, f"{folder} Profile.json")
    summ_path = os.path.join(quick, f"{folder} Summary.json")
    if not (os.path.exists(prof_path) and os.path.exists(summ_path)):
        print(f"  {folder}: no profile/summary, skipped")
        return
    # Resumable: skip issues already processed (the three fields are written
    # together, so their presence marks completion — even when blank).
    if not force and not dry:
        done = json.load(open(summ_path, encoding='utf-8'))
        if 'stakeholder_key_debates' in done:
            return
    profile = json.load(open(prof_path, encoding='utf-8'))
    sd = profile.get('source_documents') or []

    # Stakeholder Key Debates — AI over market-participant comment documents.
    sh = []
    for doc in sd:
        fname = (doc.get('file') if isinstance(doc, dict) else doc) or ""
        if classify(fname) != 'stakeholder':
            continue
        text = read_doc_text(os.path.join(base, fname))
        if not text.strip():
            continue
        sh.append((party_of(fname), (doc.get('date') if isinstance(doc, dict) else "") or "", text))

    title = profile.get('title')
    issue_id = folder
    debates = ai_summarize('stakeholder', issue_id, title, sh)
    # ERCOT / IMM Opinions — table extraction from Report docs (deterministic).
    ercot, imm = extract_opinions(base, sd)

    print(f"  {folder}: {len(sh)} participant comment doc(s)")
    print(f"     debates : {debates[:80] or '(blank)'}")
    print(f"     ercot   : {ercot[:80] or '(blank)'}")
    print(f"     imm     : {imm[:80] or '(blank)'}")
    if dry:
        return

    summary = json.load(open(summ_path, encoding='utf-8'))
    summary['stakeholder_key_debates'] = debates
    summary['ercot_opinion'] = ercot
    summary['imm_opinion'] = imm
    summary.pop('ercot_imm_opinions', None)   # retired field
    with open(summ_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)


def has_source(cat, folder):
    """True if the issue has anything to extract from: a participant comment
    (debates) or a Report doc (opinions). Others are left untouched -> 'n/a'."""
    prof = os.path.join(ROOT, cat, folder, "Quick runs", f"{folder} Profile.json")
    if not os.path.exists(prof):
        return False
    sd = json.load(open(prof, encoding='utf-8')).get('source_documents') or []
    for d in sd:
        f = (d.get('file') if isinstance(d, dict) else d) or ""
        if classify(f) == 'stakeholder':
            return True
        if 'report' in f.lower() and f.lower().endswith(('.docx', '.doc')):
            return True
    return False


def _all_folders(cat, base):
    return sorted(
        [d for d in os.listdir(base)
         if os.path.isdir(os.path.join(base, d)) and re.match(rf'{cat}\d+$', d)],
        key=lambda x: int(re.search(r'\d+', x).group()), reverse=True)


def main():
    args = sys.argv[1:]
    if not args:
        print("usage: gen_stakeholder_sections.py CAT [nums... | --recent N | --all] [--dry]")
        return
    cat = args[0].upper()
    dry = '--dry' in args
    force = '--force' in args
    base = os.path.join(ROOT, cat)

    if '--all' in args:
        picked = [f for f in _all_folders(cat, base) if has_source(cat, f)]
    elif '--recent' in args:
        n = int(args[args.index('--recent') + 1])
        picked = [f for f in _all_folders(cat, base) if has_source(cat, f)][:n]
    else:
        nums = [a for a in args[1:] if a.isdigit()]
        picked = [f"{cat}{int(x):03d}" if cat == 'COPMGRR' else f"{cat}{int(x)}" for x in nums]

    print(f"{cat}: processing {len(picked)} issue(s){' (dry run)' if dry else ''}\n")
    try:
        for folder in picked:
            process_issue(cat, folder, dry=dry, force=force)
    finally:
        quit_word()
    print(f"\nAI: {_ai['calls']} calls, {_ai['input']:,} in / {_ai['output']:,} out tokens.")


if __name__ == "__main__":
    main()
