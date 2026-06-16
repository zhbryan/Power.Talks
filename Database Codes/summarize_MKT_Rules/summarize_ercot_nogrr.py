#!/usr/bin/env python3
"""
ERCOT NOGRR Summary Creator
Implements the ERCOT Market Rules Summarization and Timeline skill for all NOGRR issues.
Reads all documents in each NOGRR folder, builds a 5-section .docx report and a
companion .json, saves to <ISSUE_ID>/Quick runs/<ISSUE_ID> Summary.docx/.json

Requires: python-docx, pywin32, openpyxl, anthropic
"""

import os, re, json
from datetime import datetime

def sanitize(text):
    if not text:
        return text
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

from docx import Document as DocxDoc
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client
import anthropic

CATEGORY = "NOGRR"

# ─── SETTINGS ────────────────────────────────────────────────────────────────
BASE_DIR      = r"E:\wamp64\www\Power.Talks\Documents Database\ERCOT.MKT.RULES\NOGRR"
AI_MODEL      = "claude-haiku-4-5-20251001"
AI_MAX_TOKENS = 600

# Haiku 4.5 pricing (USD per million tokens) — used only for the cost report.
AI_PRICE_IN   = 1.00
AI_PRICE_OUT  = 5.00

# Cumulative API token usage for the run, reported at the end of main().
_ai_usage = {"input_tokens": 0, "output_tokens": 0, "calls": 0}

# Only rebuild summaries whose issue folder gained documents since the last
# run, or whose Profile.json was refreshed after the summary was written.
# Set False to force a full rebuild of every issue.
ONLY_STALE = True

SUMMARY_DOC_EXTS = ('.pdf', '.doc', '.docx', '.xls', '.xlsx')

def issue_docs(folder):
    import os
    return sorted(f for f in os.listdir(folder)
                  if os.path.isfile(os.path.join(folder, f))
                  and os.path.splitext(f)[1].lower() in SUMMARY_DOC_EXTS)

def summary_is_stale(folder, issue_id):
    import os, json
    quick = os.path.join(folder, "Quick runs")
    jpath = os.path.join(quick, f"{issue_id} Summary.json")
    dpath = os.path.join(quick, f"{issue_id} Summary.docx")
    if not (os.path.exists(jpath) and os.path.exists(dpath)):
        return True
    try:
        with open(jpath, encoding='utf-8') as f:
            sj = json.load(f)
    except Exception:
        return True
    src = sj.get('source_documents')
    if isinstance(src, list):
        if sorted(src) != issue_docs(folder):
            return True
    else:
        dmt = os.path.getmtime(dpath)
        if any(os.path.getmtime(os.path.join(folder, f)) > dmt
               for f in issue_docs(folder)):
            return True
    ppath = os.path.join(quick, f"{issue_id} Profile.json")
    if os.path.exists(ppath) and os.path.getmtime(ppath) > os.path.getmtime(jpath):
        return True
    return False


# ─── TEXT EXTRACTION ─────────────────────────────────────────────────────────
def text_from_docx(path):
    try:
        doc = DocxDoc(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(' | '.join(c.text.strip() for c in row.cells if c.text.strip()))
        return '\n'.join(parts)
    except Exception as e:
        return f"[Error reading {os.path.basename(path)}: {e}]"

_word_app = None

def get_word():
    global _word_app
    if _word_app is None:
        _word_app = win32com.client.Dispatch("Word.Application")
        _word_app.Visible = False
    return _word_app

def quit_word():
    global _word_app
    if _word_app is not None:
        try:
            _word_app.Quit()
        except Exception:
            pass
        _word_app = None

def text_from_doc(path):
    try:
        word = get_word()
        d = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
        text = d.Content.Text
        d.Close(False)
        return text
    except Exception as e:
        return f"[Error reading {os.path.basename(path)}: {e}]"

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        return text_from_docx(path)
    elif ext == '.doc':
        return text_from_doc(path)
    return ''

# ─── BALLOT VOTE EXTRACTION ──────────────────────────────────────────────────
def extract_votes(xls_path):
    try:
        with open(xls_path, 'rb') as f:
            data = f.read()
        strings = [s.decode('ascii', errors='ignore')
                   for s in re.findall(rb'[\x20-\x7e]{8,}', data)]
        text = ' '.join(strings)
        for_m     = re.search(r'(\d+)\s*(?:For|Aye)\b', text, re.IGNORECASE)
        against_m = re.search(r'(\d+)\s*(?:Against|Nay|No)\b', text, re.IGNORECASE)
        abstain_m = re.search(r'(\d+)\s*Abstain', text, re.IGNORECASE)
        if for_m or against_m:
            f_val  = int(for_m.group(1)) if for_m else 0
            a_val  = int(against_m.group(1)) if against_m else 0
            ab_val = int(abstain_m.group(1)) if abstain_m else 0
            passed = "Passed" if f_val > a_val else "Failed"
            parts  = [f"{f_val} For", f"{a_val} Against"]
            if ab_val:
                parts.append(f"{ab_val} Abstain")
            return f"{', '.join(parts)} — {passed}"
        if re.search(r'\bpassed\b', text, re.IGNORECASE):
            return "Motion Passed"
        if re.search(r'\b(failed|defeated|withdrawn)\b', text, re.IGNORECASE):
            return "Motion Failed/Withdrawn"
    except Exception:
        pass
    return "Vote recorded"

# ─── DATE FROM FILENAME ──────────────────────────────────────────────────────
def date_from_fname(fname):
    m = re.search(r'[_\-](\d{6}|\d{8})(?:\.|$)', fname)
    if m:
        d = m.group(1)
        try:
            if len(d) == 6:
                return datetime.strptime(d, '%m%d%y').strftime('%Y-%m-%d')
            else:
                return datetime.strptime(d, '%m%d%Y').strftime('%Y-%m-%d')
        except ValueError:
            pass
    return None

# ─── TIMELINE BUILDER ────────────────────────────────────────────────────────
_BALLOT_PAT = re.compile(r'(ros|tac|prs|rms|wms|cops)[\s_-]*(?:email[\s_-]*)?ballot', re.IGNORECASE)
_REPORT_PAT = re.compile(r'(ros|tac|prs|rms|wms|cops)[\s_-]*report', re.IGNORECASE)
_BOARD_PAT  = re.compile(r'board.?report', re.IGNORECASE)
_PUCT_PAT   = re.compile(r'puct.?report', re.IGNORECASE)
_IMPACT_PAT = re.compile(r'impact.?analysis', re.IGNORECASE)

def build_timeline(folder):
    events = []
    for fname in os.listdir(folder):
        fl    = fname.lower()
        fpath = os.path.join(folder, fname)
        date_str = date_from_fname(fname) or 'Unknown'
        if _BALLOT_PAT.search(fl):
            body = _BALLOT_PAT.search(fl).group(1).upper()
            events.append({'date': date_str, 'body': body, 'action': 'Ballot', 'outcome': extract_votes(fpath)})
        elif _BOARD_PAT.search(fl):
            events.append({'date': date_str, 'body': 'Board', 'action': 'Board Report', 'outcome': 'Board action'})
        elif _PUCT_PAT.search(fl):
            events.append({'date': date_str, 'body': 'PUCT', 'action': 'PUCT Filing', 'outcome': 'Regulatory filing'})
        elif _REPORT_PAT.search(fl):
            body = _REPORT_PAT.search(fl).group(1).upper()
            events.append({'date': date_str, 'body': body, 'action': 'Committee Report', 'outcome': 'See report'})
    events.sort(key=lambda x: x['date'] if x['date'] != 'Unknown' else '9999-99-99')
    return events

# ─── STATUS INFERENCE ────────────────────────────────────────────────────────
def infer_status(folder):
    files = [f.lower() for f in os.listdir(folder)]
    if any('withdrawal' in f for f in files):
        return "Withdrawn"
    if any(_BOARD_PAT.search(f) for f in files) or any(_PUCT_PAT.search(f) for f in files):
        return "Approved"
    if any(_REPORT_PAT.search(f) and 'tac' in f for f in files):
        return "Pending Board Action"
    if any(_REPORT_PAT.search(f) and ('ros' in f or 'prs' in f or 'rms' in f or 'wms' in f or 'cops' in f) for f in files):
        return "In ROS/TAC Review"
    return "Recently Posted / Pending ROS"

# ─── AI EXECUTIVE SUMMARY ────────────────────────────────────────────────────
_ai_client = None

def get_ai():
    global _ai_client
    if _ai_client is None:
        from anthropic_key import get_anthropic_key
        _ai_client = anthropic.Anthropic(api_key=get_anthropic_key())
    return _ai_client

def ai_executive_summary(issue_id, title, revision_desc, reason, business_case, sections, status):
    prompt = f"""Write a concise executive summary (3–5 sentences, plain English, no jargon) for this ERCOT Nodal Operating Guide Revision Request (NOGRR).

Issue: {issue_id}
Title: {title or 'N/A'}
Nodal Operating Guide Sections: {', '.join(sections) if sections else 'Not specified'}
Current Status: {status}

Revision Description:
{revision_desc or 'Not provided'}

Reason for Revision:
{reason or 'Not provided'}

Business Case / Justification:
{business_case or 'Not provided'}

Write 3–5 sentences covering: (1) what operating procedure or guide section is changing, (2) why the change is needed, (3) who is affected and how. Note how this differs from a protocol-level change where relevant. Do not use markdown. Do not repeat the issue number in the summary."""

    try:
        msg = get_ai().messages.create(
            model=AI_MODEL, max_tokens=AI_MAX_TOKENS,
            messages=[{"role": "user", "content": prompt}]
        )
        _ai_usage["input_tokens"]  += msg.usage.input_tokens
        _ai_usage["output_tokens"] += msg.usage.output_tokens
        _ai_usage["calls"]         += 1
        return msg.content[0].text.strip()
    except Exception:
        return (f"This Nodal Operating Guide Revision Request ({issue_id}) proposes changes to "
                f"{', '.join(sections) or 'Nodal Operating Guide sections'}. "
                f"{revision_desc or ''}")

# ─── DOCX REPORT WRITER ──────────────────────────────────────────────────────
def write_summary_docx(out_path, issue_id, title, exec_summary, issue_details,
                        impact_items, timeline, current_status):
    doc = DocxDoc()
    h = doc.add_heading(f"{issue_id} — {title or 'Nodal Operating Guide Revision Request'}", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Summary Report")
    doc.add_paragraph()

    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(sanitize(exec_summary))

    doc.add_heading("2. Issue Details", level=2)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    for label, value in issue_details:
        row = tbl.add_row()
        row.cells[0].text = sanitize(label) or label
        row.cells[1].text = sanitize(str(value)) if value else '—'
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("3. Impact Analysis", level=2)
    if impact_items:
        for _lbl, _val in impact_items:
            doc.add_paragraph(f"{sanitize(_lbl)}: {sanitize(_val)}", style="List Bullet")
    else:
        doc.add_paragraph("Impact analysis document not available for this issue.")

    doc.add_heading("4. Stakeholder Discussion Timeline", level=2)
    if timeline:
        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.style = 'Table Grid'
        hdr = tbl2.rows[0].cells
        for i, h in enumerate(["Date", "Body", "Action", "Outcome"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for ev in timeline:
            row = tbl2.add_row()
            row.cells[0].text = ev['date']
            row.cells[1].text = ev['body']
            row.cells[2].text = ev['action']
            row.cells[3].text = ev['outcome']
    else:
        doc.add_paragraph("No ballot or committee report documents found for this issue.")

    doc.add_heading("5. Current Status", level=2)
    doc.add_paragraph(current_status)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)

# ─── HELPERS ─────────────────────────────────────────────────────────────────
def load_profile(folder, issue_id):
    path = os.path.join(folder, "Quick runs", f"{issue_id} Profile.json")
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def extract_impact_items(folder):
    """Parse the most recent impact analysis document into (label, value)
    bullet items. Table rows become labeled bullets; plain text becomes
    Note bullets as a fallback."""
    ia_files = [f for f in sorted(os.listdir(folder))
                if _IMPACT_PAT.search(f.lower())
                and os.path.splitext(f)[1].lower() in ('.docx', '.doc')]
    if not ia_files:
        return []
    path = os.path.join(folder, ia_files[-1])  # most recent revision
    items = []
    seen = set()

    def add(label, value):
        label = ' '.join(label.split())
        value = ' '.join(value.split())
        if not label or not value or len(label) > 90 or label.lower() == value.lower():
            return
        key = label.lower()
        if re.fullmatch(r"(nprr|nogrr|pgrr|rmgrr|scr|copmgrr)\s*(number|title)", key):
            return
        if key in seen:
            return
        seen.add(key)
        items.append((label, value[:400]))

    try:
        if path.lower().endswith('.docx'):
            d = DocxDoc(path)
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells]
                    uniq = list(dict.fromkeys(c for c in cells if c))
                    if len(uniq) >= 2:
                        add(uniq[0], uniq[1])
        else:
            word = get_word()
            wd = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
            try:
                for tbl in wd.Tables:
                    try:
                        n_rows, n_cols = tbl.Rows.Count, tbl.Columns.Count
                    except Exception:
                        continue
                    for r in range(1, n_rows + 1):
                        cells = []
                        for c in range(1, min(n_cols, 3) + 1):
                            try:
                                cells.append(tbl.Cell(r, c).Range.Text.replace('\r', ' ').replace('\x07', '').strip())
                            except Exception:
                                continue
                        uniq = list(dict.fromkeys(c for c in cells if c))
                        if len(uniq) >= 2:
                            add(uniq[0], uniq[1])
            finally:
                wd.Close(False)
    except Exception:
        pass

    if not items:
        text = extract_text(path)[:2000].strip()
        for line in re.split(r'[\r\n]+', text):
            line = line.strip(' \t-')
            if 15 <= len(line) <= 300:
                items.append(('Note', line))
            if len(items) >= 8:
                break
    return items[:12]

def status_sentence(issue_id, inferred):
    if inferred == "Approved":
        return f"{issue_id} has been approved by the ERCOT Board of Directors."
    if inferred == "Withdrawn":
        return f"{issue_id} has been withdrawn by the sponsor."
    if inferred == "Pending Board Action":
        return f"{issue_id} has passed TAC and is pending Board of Directors action."
    if inferred == "In ROS/TAC Review":
        return f"{issue_id} is currently under review by the ROS or TAC committee."
    return f"{issue_id} has been recently posted and is pending initial committee review."

# ─── PER-ISSUE PROCESSOR ─────────────────────────────────────────────────────
def process_issue(folder, n):
    issue_id  = f"NOGRR{n}"
    quick     = os.path.join(folder, "Quick runs")
    out_path  = os.path.join(quick, f"{issue_id} Summary.docx")

    profile   = load_profile(folder, issue_id)
    title     = profile.get('title') or issue_id
    sections  = profile.get('governing_document_sections', [])
    rev_desc  = profile.get('revision_description')
    _reason   = profile.get('reason_for_revision')
    reason    = '; '.join(_reason) if isinstance(_reason, list) else (_reason or '')
    biz_case  = profile.get('business_case')
    status    = profile.get('status') or infer_status(folder)
    date_post = profile.get('date_posted_decision')
    req_res   = profile.get('timeline_requested_resolution')
    sponsor   = profile.get('sponsor_name')
    company   = profile.get('sponsor_company')
    mkt_seg   = profile.get('market_segment')

    exec_summary = ai_executive_summary(issue_id, title, rev_desc, reason, biz_case, sections, status)

    issue_details = [
        ("Status",                                    status),
        ("Date Posted",                               date_post or '—'),
        ("Requested Resolution",                      req_res or '—'),
        ("Nodal Operating Guide Sections Revised",    '\n'.join(sections) if sections else '—'),
        ("Sponsor",                                   f"{sponsor} ({company})" if sponsor else '—'),
        ("Market Segment",                            mkt_seg or '—'),
    ]

    impact_items    = extract_impact_items(folder)
    timeline       = build_timeline(folder)
    inferred       = infer_status(folder)
    current_status = status_sentence(issue_id, inferred)

    write_summary_docx(out_path, issue_id, title, exec_summary, issue_details,
                        impact_items, timeline, current_status)

    sponsor_str = f"{sponsor} · {company}" if sponsor and company else (sponsor or '')
    summary_json = {
        "source_documents":   issue_docs(folder),
        "summary_last_updated": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
        "nogrr_number":       n,
        "title":              title,
        "status":             status,
        "date_posted":        date_post or "",
        "effective_date":     profile.get("effective_date") or "",
        "guide_sections":     sections,
        "sponsor":            sponsor_str,
        "executive_summary":  exec_summary,
        "background":         reason,
        "key_change":         rev_desc or "",
        "impacts":            [{"category": "Business Case / Justification", "text": biz_case[:500]}] if biz_case else [],
        "impact_analysis":   [{"label": "Impact Analysis", "rows": [[l, v] for l, v in impact_items]}] if impact_items else [],
        "timeline":           [{"date": ev["date"], "body": ev["body"], "action": ev["action"], "notes": ev["outcome"]} for ev in timeline],
        "current_status":     [current_status],
    }
    with open(os.path.join(quick, f"{issue_id} Summary.json"), 'w', encoding='utf-8') as f:
        json.dump(summary_json, f, indent=2, ensure_ascii=False)

    return out_path

# ─── MAIN ────────────────────────────────────────────────────────────────────
def main():
    folders = sorted([
        d for d in os.listdir(BASE_DIR)
        if os.path.isdir(os.path.join(BASE_DIR, d)) and re.match(r'NOGRR\d+', d)
    ], key=lambda x: int(re.search(r'\d+', x).group()))

    if ONLY_STALE:
        folders = [d for d in folders
                   if summary_is_stale(os.path.join(BASE_DIR, d), d)]
        print(f"ONLY_STALE: {len(folders)} issue(s) need a summary rebuild.")

    print(f"Processing {len(folders)} NOGRR folders...\n")
    ok = err = 0
    try:
        for folder_name in folders:
            n      = int(re.search(r'\d+', folder_name).group())
            folder = os.path.join(BASE_DIR, folder_name)
            try:
                out = process_issue(folder, n)
                print(f"[OK] NOGRR{n:>4}  ->  {os.path.basename(out)}")
                ok += 1
            except Exception as e:
                print(f"[ERR] NOGRR{n}: {e}")
                err += 1
    finally:
        quit_word()
    print(f"\nDone: {ok} summaries created, {err} errors.")

    in_tok  = _ai_usage["input_tokens"]
    out_tok = _ai_usage["output_tokens"]
    cost_in  = in_tok  / 1_000_000 * AI_PRICE_IN
    cost_out = out_tok / 1_000_000 * AI_PRICE_OUT
    print(f"\nAI usage ({_ai_usage['calls']} calls, {AI_MODEL}):")
    print(f"  input : {in_tok:>8,} tok  ->  ${cost_in:.4f}")
    print(f"  output: {out_tok:>8,} tok  ->  ${cost_out:.4f}")
    print(f"  TOTAL : ${cost_in + cost_out:.4f}")

if __name__ == "__main__":
    main()
