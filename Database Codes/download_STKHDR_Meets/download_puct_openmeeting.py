#!/usr/bin/env python3
"""
Download a PUCT Open Meeting recording from AdminMonitor.

AdminMonitor posts each PUCT open meeting as an on-demand HLS video (VOD) a
short time after the meeting ends. This is far more reliable than trying to
live-capture ERCOT's stream at an exact scheduled time (the old approach, which
also failed on an unquoted Windows path). Given a meeting date, this:

  1. fetches the meeting page   /tx/puct/open_meeting/<YYYYMMDD>/
  2. extracts the <source src="…master.m3u8"> (the CloudFront VOD playlist)
  3. downloads it with ffmpeg to MP3 (audio, default) and/or MP4 (video)
  4. fetches the FINAL agenda -> .docx (the archived, status_id=3 agenda linked
     on the meeting page = the agenda as of the meeting's END, not the live
     start-time one)

The stream id (e.g. am_3473_20260820_PUCT_OM_A1) is meeting-specific and NOT
derivable from the date, so the page MUST be scraped for the real URL.

Output (Power.Talks convention: downloaded media -> Documents Database/):
  Documents Database/ERCOT.STKHDR.MEETS/PUCT/PUCT_OpenMeeting_<YYYYMMDD>.mp3
  Documents Database/ERCOT.STKHDR.MEETS/PUCT/PUCT_OpenMeeting_<YYYYMMDD>.mp4  (with --video)
  Documents Database/ERCOT.STKHDR.MEETS/PUCT/PUCT_Agenda_<YYYYMMDD>.docx        (unless --no-agenda)

Usage:
  py -3 download_puct_openmeeting.py 20260820           # audio MP3 + final agenda
  py -3 download_puct_openmeeting.py 20260820 --video    # also full MP4
  py -3 download_puct_openmeeting.py 20260820 --video-only
  py -3 download_puct_openmeeting.py 20260820 --no-agenda
"""

import argparse
import os
import re
import shutil
import subprocess
import sys

import requests

UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
      "(KHTML, like Gecko) Chrome/128.0.0.0 Safari/537.36")
ADMIN_HOST = "https://www.adminmonitor.com"
BASE = f"{ADMIN_HOST}/tx/puct/open_meeting"

HERE = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
OUT_DIR = os.path.join(PROJECT_ROOT, "Documents Database",
                       "ERCOT.STKHDR.MEETS", "PUCT")

# ffmpeg installed via winget lands here; fall back to PATH.
FFMPEG = (r"C:\Users\chunl\AppData\Local\Microsoft\WinGet\Links\ffmpeg.exe"
          if os.path.exists(r"C:\Users\chunl\AppData\Local\Microsoft\WinGet\Links\ffmpeg.exe")
          else shutil.which("ffmpeg") or "ffmpeg")


def get_meeting_html(date_yyyymmdd):
    """Fetch the meeting page HTML, or None if it doesn't exist (404).
    Raises on other HTTP errors. Fetch once, then parse for both video + agenda."""
    url = f"{BASE}/{date_yyyymmdd}/"
    r = requests.get(url, headers={"User-Agent": UA, "Referer": BASE + "/"},
                     timeout=60)
    if r.status_code == 404:
        return None
    if r.status_code != 200:
        raise RuntimeError(f"meeting page {url} -> HTTP {r.status_code}")
    return r.text


def master_m3u8_from_html(html):
    """The HLS master playlist URL from a meeting page, or None if no VOD yet."""
    m = re.search(r'<source[^>]+src="([^"]+\.m3u8)"', html, re.IGNORECASE)
    return m.group(1) if m else None


def agenda_url_from_html(html):
    """The FINAL agenda iframe URL from a meeting page, made absolute, or None.

    Only the archived agenda is returned. AdminMonitor serves two variants:
      * before/at the meeting: .../live/inc_agenda_text.cfm?...&status_id=   (start-time)
      * after the meeting ends: .../archive/inc_agenda_text.cfm?...&status_id=3 (final)
    We want the meeting's END-state agenda, so a page still showing only the
    live/start-time agenda (meeting not yet concluded) yields None — retry later."""
    m = re.search(r'src="([^"]*inc_agenda_text\.cfm[^"]*)"', html, re.IGNORECASE)
    if not m:
        return None
    url = m.group(1)
    if "/archive/" not in url.lower():
        return None   # still the live/start-time agenda; meeting not ended yet
    if url.startswith("/"):
        url = ADMIN_HOST + url
    return url


def find_master_m3u8(date_yyyymmdd):
    """Convenience: fetch the page and return its HLS master playlist URL (or
    None if the meeting page/VOD isn't posted yet)."""
    html = get_meeting_html(date_yyyymmdd)
    return master_m3u8_from_html(html) if html else None


def ffmpeg_pull(m3u8, out_path, audio_only):
    """Download the HLS VOD to out_path. Audio -> MP3; video -> MP4 (copy).
    Raises RuntimeError on failure."""
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    if audio_only:
        args = ["-vn", "-acodec", "libmp3lame", "-q:a", "4"]
    else:
        # Pick the highest-resolution variant, stream-copy (fast, lossless).
        args = ["-map", "0:p:0:v:0?", "-map", "0:p:0:a:0?",
                "-c", "copy", "-bsf:a", "aac_adtstoasc"]
    cmd = [FFMPEG, "-y", "-i", m3u8, *args, out_path]
    print("  " + " ".join(cmd))
    rc = subprocess.call(cmd)
    if rc != 0 or not os.path.exists(out_path):
        raise RuntimeError(f"ffmpeg failed (rc={rc}) for {out_path}")
    mb = os.path.getsize(out_path) / (1024 * 1024)
    print(f"  saved {out_path} ({mb:.1f} MB)")


def fetch_agenda_docx(agenda_url, out_path):
    """Download the (final/archived) agenda HTML and write it to a .docx.
    Raises RuntimeError on failure."""
    from bs4 import BeautifulSoup
    from docx import Document

    r = requests.get(agenda_url, headers={"User-Agent": UA, "Referer": BASE + "/"},
                     timeout=60)
    if r.status_code != 200:
        raise RuntimeError(f"agenda {agenda_url} -> HTTP {r.status_code}")
    soup = BeautifulSoup(r.text, "html.parser")

    doc = Document()
    h2 = soup.find("h2")
    if h2:
        doc.add_heading(h2.get_text(strip=True), level=1)
    h3 = soup.find("h3")
    if h3:
        doc.add_heading(h3.get_text(" ", strip=True), level=2)
    for p in soup.find_all("p"):
        if p.find_parent("table"):
            continue
        text = p.get_text(" ", strip=True)
        if text:
            doc.add_paragraph(text)
    table = soup.find("table")
    if table:
        doc.add_heading("Agenda Items", level=2)
        for tr in table.find_all("tr"):
            cells = [td.get_text(" ", strip=True) for td in tr.find_all("td")]
            line = " ".join(c for c in cells if c)
            if line:
                doc.add_paragraph(line)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"  saved {out_path}")


def main():
    ap = argparse.ArgumentParser(description="Download a PUCT open meeting recording + final agenda.")
    ap.add_argument("date", help="meeting date, YYYYMMDD (e.g. 20260820)")
    ap.add_argument("--video", action="store_true", help="also download full MP4")
    ap.add_argument("--video-only", action="store_true", help="download MP4 only")
    ap.add_argument("--no-agenda", action="store_true",
                    help="skip fetching the final (archived) agenda .docx")
    args = ap.parse_args()

    if not re.fullmatch(r"\d{8}", args.date):
        sys.exit("ERROR: date must be YYYYMMDD, e.g. 20260820")

    print(f"Locating PUCT open meeting for {args.date}…")
    try:
        html = get_meeting_html(args.date)
    except RuntimeError as e:
        sys.exit(f"ERROR: {e}")
    if not html:
        sys.exit(f"ERROR: no meeting page for {args.date}.")

    base = os.path.join(OUT_DIR, f"PUCT_OpenMeeting_{args.date}")
    try:
        m3u8 = master_m3u8_from_html(html)
        if not m3u8:
            sys.exit(f"ERROR: no VOD found for {args.date} (meeting not posted yet?).")
        print(f"  stream: {m3u8}")
        if not args.video_only:
            ffmpeg_pull(m3u8, base + ".mp3", audio_only=True)
        if args.video or args.video_only:
            ffmpeg_pull(m3u8, base + ".mp4", audio_only=False)

        # Final agenda (archived, status_id=3 = as of the meeting's end).
        if not args.no_agenda:
            agenda_url = agenda_url_from_html(html)
            if agenda_url:
                fetch_agenda_docx(agenda_url, os.path.join(
                    OUT_DIR, f"PUCT_Agenda_{args.date}.docx"))
            else:
                print("  (no agenda iframe found on page)")
    except RuntimeError as e:
        sys.exit(f"ERROR: {e}")
    print("Done.")


if __name__ == "__main__":
    main()
