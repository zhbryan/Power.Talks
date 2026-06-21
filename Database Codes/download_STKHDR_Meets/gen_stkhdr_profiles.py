#!/usr/bin/env python3
"""Batch-generate ERCOT Stakeholder Meeting Profile JSONs for every meeting.

Implements the `ERCOT Stakeholder Meeting Profile` skill at scale: for each
committee/date folder under Documents Database/ERCOT.STKHDR.MEETS, write
<COMMITTEE>/<DATE>/Quick runs/<COMMITTEE>-<DATE> Profile.json.

Rules honored from the skill:
  - documents EXCLUDES .zip archives (case-insensitive), .tmp files, _manifest.json
  - group_summary carries leadership + voting_parties + voting_structure (per group)
  - empty arrays are [], unknown scalars are null
Existing profiles are left untouched (re-runnable; preserves curated ones).

Per-meeting extraction is best-effort:
  - agenda_items   from the *agenda* .docx (tables + numbered paragraphs)
  - ballot_results from the *ballot* .xls/.xlsx (motion lines; tallies left null)
  - working_group_reports inferred from report/update filenames
Old binary .doc agendas are not parsed (listed in documents only).
"""
import os, re, json, sys

ROOT = os.path.join("Documents Database", "ERCOT.STKHDR.MEETS")
DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")

# ── ERCOT market segments (subcommittee/TAC voting) ──────────────────────────
SEG_SUB = [
    "Consumer (Residential / Commercial / Industrial subsegments)",
    "Cooperative", "Independent Generator", "Independent Power Marketer",
    "Independent Retail Electric Provider", "Investor Owned Utility", "Municipal",
]

# Working-group abbreviations used to detect WG reports in filenames.
WG_ABBRS = ["BSWG","DWG","IBRWG","MWG","NDSWG","OWG","OTWG","PDCWG","PLWG",
            "SPWG","SSWG","VPWG","CMWG","DSWG","SAWG","WMWG","TDTMS","RMTTF"]

# Agenda boilerplate to drop from key_discussion_topics.
TOPIC_SKIP = re.compile(r"antitrust|agenda review|adjourn|other business|"
                        r"approval of .*minutes|roll call|opening remarks", re.I)

# ── Committee metadata (leadership + voting model), from the Links skill §1 ───
# model: "board" (governance), "segment" (TAC + subcommittees), "wg" (advisory)
C = {
 "BOARD": ("Board of Directors", "Bill Flores", "Peggy Heeg", "board", None),
 "FA":    ("Finance and Audit Committee", "Christopher Krummel", None, "board", "Board"),
 "HRG":   ("HR and Governance Committee", "Peggy Heeg", None, "board", "Board"),
 "TS":    ("Technology and Security Committee", "John Swainson", None, "board", "Board"),
 "TAC":   ("Technical Advisory Committee", "Caitlin Smith", "Martha Henson", "segment", "Board"),
 "PRS":   ("Protocol Revision Subcommittee", "Diana Coleman", "Eric Blakey", "segment", "TAC"),
 "ROS":   ("Reliability and Operations Subcommittee", "Sandeep Borkar", "Shane Thomas", "segment", "TAC"),
 "RMS":   ("Retail Market Subcommittee", "John Schatz", "Debbie McKeever", "segment", "TAC"),
 "WMS":   ("Wholesale Market Subcommittee", "Blake Holt", "Jim Lee", "segment", "TAC"),
 "LLWG":  ("Large Load Working Group", "Bob Wittmeyer", "Patrick Gravois", "wg", "TAC"),
 "CFSG":  ("Credit Finance Sub Group", "Jett Price", "Loretto Martin", "wg", "TAC"),
 "RTCBTF":("Real-Time Co-optimization plus Batteries Task Force", None, None, "wg", "TAC"),
 "BESTF": ("Battery Energy Storage Task Force", None, None, "wg", "TAC"),
 "TDTMS": ("Texas Data Transport and MarkeTrak Systems Working Group", "Sheri Wiegand", "Rob Bevill / Monica Jones", "wg", "RMS"),
 "RMTTF": ("Retail Market Training Task Force", "Melinda Earnest / Tomas Fernandez / Deborah McKeever", None, "wg", "RMS"),
 "TXSET": ("Texas Standard Electronic Transaction (Texas SET) Working Group", None, None, "wg", "RMS"),
 "TXSETLP":("Texas SET Lite Process Working Group", "Sam Pak / Kyle Patrick / Stephen Wilson (Co-Chairs)", None, "wg", "RMS"),
 "BSWG":  ("Black Start Working Group", "Michael Dieringer", "Cerina Rivera-Terrell", "wg", "ROS"),
 "DWG":   ("Dynamics Working Group", "Aditi Upadhyay", "Xuan Wu", "wg", "ROS"),
 "IBRWG": ("Inverter-Based Resource Working Group", "Julia Matevosyan", "Miguel Cova Acosta", "wg", "ROS"),
 "MWG":   ("Meter Working Group", "Kyle Stuckly", "Tony Davis", "wg", "ROS"),
 "NDSWG": ("Network Data Support Working Group", "Teddi Flessner", "Vincent Cutlip", "wg", "ROS"),
 "OWG":   ("Operations Working Group", "Rickey Floyd", "Tyler Springer", "wg", "ROS"),
 "OTWG":  ("Operations Training Working Group", "Manuel Sanchez", "Mike Flores", "wg", "ROS"),
 "PDCWG": ("Performance, Disturbance, Compliance Working Group", "Paul Messmann", "Luis Hinojosa", "wg", "ROS"),
 "PLWG":  ("Planning Working Group", "Mina Turner", "Kiran Kota", "wg", "ROS"),
 "SPWG":  ("System Protection Working Group", "Uchenna Ndusorouwa", "Jourdan Watkins", "wg", "ROS"),
 "SSWG":  ("Steady State Working Group", "Chris Ramirez", "Weiwei Hu", "wg", "ROS"),
 "VPWG":  ("Voltage Profile Working Group", "Charles Aleman", "Alison Flint", "wg", "ROS"),
 "CMWG":  ("Congestion Management Working Group", "Chenyan Guo", "Shane Thomas", "wg", "WMS"),
 "DSWG":  ("Demand Side Working Group", "Nathaniel Mancha", None, "wg", "WMS"),
 "SAWG":  ("Supply Analysis Working Group", "Kevin Hanson", "Pete Warnken / Greg Lackey", "wg", "WMS"),
 "WMWG":  ("Wholesale Market Working Group", "Amanda Frazier", "Trevor Safko", "wg", "WMS"),
}

CAL_SLUG = {  # calendar URL slug per committee (best-effort, base form)
 "BOARD":"Board-of-Directors-Meeting", "FA":"Finance-_-Audit-Committee",
 "HRG":"HR-_-Governance-Committee", "TS":"Technology-_-Security-Committee",
 "TAC":"TAC-Meeting", "TXSET":"Texas-SET_LP-Meeting", "TXSETLP":"Texas-SET_LP-Meeting",
}


# ── ERCOT Board of Directors — voting members (the "voting parties") ─────────
# UPDATE SCHEME: the Board is not market-segment-based, so its voting parties are
# the named directors. Refresh this registry from the live ERCOT pages whenever
# the Board changes (ERCOT posts a news release when directors are appointed):
#   • Independent directors:  https://www.ercot.com/about/governance/directors
#   • Committee rosters:      https://www.ercot.com/committees/board/finance_audit
#                             https://www.ercot.com/committees/board/hr_governance
#                             https://www.ercot.com/committees/board/tech-security
# After editing, bump BOARD_LAST_VERIFIED and regenerate the manifests
# (gen_stkhdr_manifest.py) so the group homepage's Voting Parties list updates.
BOARD_LAST_VERIFIED = "2026-06"
BOARD_MEMBERS = {
    "BOARD": ["Bill Flores (Chair)", "Peggy Heeg (Vice Chair)", "Linda Capuano",
              "Julie England", "Christopher A. Krummel", "Kathleen McAllister",
              "Bill Mohl", "John Swainson"],
    "FA":  ["Christopher A. Krummel (Chair)", "Benjamin Barkley", "Julie England", "Bill Mohl"],
    "HRG": ["Peggy Heeg (Chair)", "Benjamin Barkley", "Christopher A. Krummel", "Kathleen McAllister"],
    "TS":  ["John Swainson (Chair)", "Linda Capuano", "Julie England", "Bill Mohl"],
}
# Ex officio (sit on the Board but the President/CEO and PUCT Chair are non-voting).
BOARD_EX_OFFICIO = ["Pablo Vegas (President & CEO, non-voting)",
                    "Thomas Gleeson (PUCT Chair, non-voting)",
                    "Courtney Hjaltman (PUCT Commissioner)",
                    "Benjamin Barkley (OPUC Public Counsel)"]


def group_summary(abbr):
    full, chair, vc, model, parent = C.get(abbr, (abbr, None, None, "wg", "its parent committee"))
    lead_bits = []
    if chair: lead_bits.append(f"Chair: {chair}")
    if vc: lead_bits.append(f"Vice Chair: {vc}")
    leadership = "; ".join(lead_bits) if lead_bits else None

    if model == "board":
        parties = BOARD_MEMBERS.get(abbr, ["ERCOT Board of Directors"])
        if abbr == "BOARD":
            overview = ("The ERCOT Board of Directors is the governing body of ERCOT, "
                        "overseeing management, finance, and the reliability of the Texas grid.")
            structure = ("The eight independent directors hold the votes; matters are decided "
                         "by a majority. Ex officio members participate but the President & CEO "
                         "(Pablo Vegas) and PUCT Chair (Thomas Gleeson) are non-voting; "
                         "Courtney Hjaltman (PUCT Commissioner) and Benjamin Barkley (OPUC "
                         f"Public Counsel) also serve ex officio. Not subject to ERCOT "
                         f"market-segment voting. Roster verified {BOARD_LAST_VERIFIED}.")
        else:
            overview = (f"The {full} is a standing committee of the ERCOT Board of Directors, "
                        f"reviewing matters in its area and forwarding recommendations to the "
                        f"full Board.")
            structure = ("Decisions by a majority of the committee's directors, forwarded to "
                         f"the full ERCOT Board. Not subject to ERCOT market-segment voting. "
                         f"Roster verified {BOARD_LAST_VERIFIED}.")
    elif model == "segment":
        if abbr == "TAC":
            overview = ("The Technical Advisory Committee (TAC) is the senior stakeholder "
                        "committee reporting to the ERCOT Board, reviewing and forwarding "
                        "recommendations from the subcommittees (PRS, ROS, RMS, WMS).")
            structure = ("ERCOT market-segment voting at the TAC level — votes are tallied "
                         "by market segment, not headcount; recommendations are forwarded "
                         "to the ERCOT Board.")
        else:
            overview = (f"The {full} ({abbr}) reports to the Technical Advisory Committee "
                        f"(TAC) and develops and votes on revision requests and "
                        f"recommendations in its market domain.")
            structure = ("ERCOT market-segment voting: votes are tallied by market segment "
                         "(not headcount) and a typical motion passes on a majority of "
                         "segment votes. The Consumer segment carries 1.5 votes for "
                         "ROS/RMS/WMS (1 for PRS), split among present subsegments. "
                         "Endorsed items are forwarded to TAC.")
        parties = list(SEG_SUB)
    else:  # wg
        overview = (f"The {full} ({abbr}) is a technical working group under {parent}, "
                    f"developing analysis and recommendations in its subject area and "
                    f"reporting up to its parent committee.")
        parties = []
        structure = (f"Advisory working group — positions and straw votes are forwarded to "
                     f"{parent}; formal balloting occurs at the subcommittee/TAC level "
                     f"under ERCOT market-segment voting.")
    return full, chair, vc, {
        "overview": overview, "leadership": leadership,
        "voting_parties": parties, "voting_structure": structure,
    }


def extract_agenda(path):
    return _agenda_pdf(path) if path.lower().endswith(".pdf") else _agenda_docx(path)


def _agenda_docx(path):
    try:
        from docx import Document
        doc = Document(path)
    except Exception:
        return []
    items, seen = [], set()
    def add(t):
        t = re.sub(r"\s+", " ", t).strip()
        if t and t.lower() not in seen and len(t) < 200:
            seen.add(t.lower()); items.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells and re.match(r"^\d+\.$", cells[0]):
                title = next((c for c in cells[1:] if c), "")
                if title and not re.match(r"^\d", title):
                    add(title)
    for p in doc.paragraphs:
        m = re.match(r"^\d+\.\s+(.+)", p.text.strip())
        if m:
            add(m.group(1))
    return items[:40]


def _agenda_pdf(path):
    """Extract numbered agenda items from a PDF, using column geometry to keep
    the Topic text and drop the Topic-Type / Presenter columns."""
    try:
        import pdfplumber
    except Exception:
        return []
    items, seen = [], set()
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                words = page.extract_words()
                if not words:
                    continue
                top_min, top_max = page.height * 0.07, page.height * 0.93
                lines = {}
                for w in words:
                    if top_min <= w["top"] <= top_max:  # drop header/footer margins
                        lines.setdefault(round(w["top"] / 3.0), []).append(w)
                cutoff = None
                for k in sorted(lines):
                    txt = [w["text"] for w in lines[k]]
                    if "Presenter" in txt and "Type" in txt:
                        tops = [w for w in lines[k] if w["text"] == "Topic"]
                        cutoff = ((tops[-1]["x0"] if tops else
                                   [w for w in lines[k] if w["text"] == "Type"][0]["x0"]) - 4)
                        break
                if cutoff is None:
                    cutoff = page.width * 0.60
                cur = None
                for k in sorted(lines):
                    ws = sorted(lines[k], key=lambda w: w["x0"])
                    line = re.sub(r"\s+", " ",
                                  " ".join(w["text"] for w in ws if w["x0"] < cutoff)).strip()
                    if re.match(r"^(\d+)\.\s+", line):
                        if cur:
                            items.append(cur)
                        cur = re.sub(r"^\d+\.\s+", "", line)
                    elif re.match(r"^\d+\.\d+", line):  # sub-item → close current
                        if cur:
                            items.append(cur); cur = None
                    elif (cur is not None and line
                          and not re.match(r"^(Item|Convene|Adjourn|Recess|Executive Session"
                                           r"|Agenda|ERCOT Public|Page)\b", line)):
                        cur += " " + line
                if cur:
                    items.append(cur)
    except Exception:
        return []
    clean = []
    for it in items:
        it = re.sub(r"\s+", " ", it).strip().rstrip(",")
        if it and it.lower() not in seen and 3 < len(it) < 160:
            seen.add(it.lower()); clean.append(it)
    return clean[:40]


_RR_CODE = re.compile(r"\b((?:NPRR|NOGRR|PGRR|RRGRR|SCR|RMGRR|COPMGRR|VCMRR|"
                      r"OBDRR|SMOGRR)\d{2,4})\b")


def _read_sheets(path):
    """Return {sheet_name: [[cell, …], …]} for .xls (xlrd) or .xlsx (openpyxl)."""
    low = path.lower()
    try:
        if low.endswith(".xlsx"):
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True, data_only=True)
            return {ws.title: [list(r) for r in ws.iter_rows(values_only=True)]
                    for ws in wb.worksheets}
        import xlrd
        bk = xlrd.open_workbook(path)
        return {sh.name: [[sh.cell_value(r, c) for c in range(sh.ncols)]
                          for r in range(sh.nrows)] for sh in bk.sheets()}
    except Exception:
        return {}


def _is_num(x):
    return bool(re.match(r"^-?\d+(?:\.\d+)?$", str(x).strip()))


def _parse_vote_sheet(rows):
    """From a ballot 'Vote' sheet read the result and the **total** tally
    (Yes / No / Abstain). Layouts differ — the total row is labelled
    '<COMMITTEE> Vote:' (TAC) or 'Segment Vote:' (PRS) and sits *above* the
    per-entity data header ('… Yes | No | Abstain'); per-segment 'Segment Vote:'
    rows sit below it. So we read the first tally row before that header.
    Returns (result, for, against, abstain)."""
    grid = [[("" if c is None else str(c).strip()) for c in row] for row in rows]

    # result — anywhere on the sheet
    result = None
    for cells in grid:
        m = re.search(r"motion\s+(passes|carries|fails)",
                      " ".join(c for c in cells if c).lower())
        if m:
            result = "Failed" if m.group(1) == "fails" else "Passed"
            break

    # locate the per-entity data header (Yes/No/Abstain columns)
    hdr = len(grid)
    for i, cells in enumerate(grid):
        low = [c.lower() for c in cells]
        if ("yes" in low and "no" in low and "abstain" in low) or \
           any("sector / entity" in c for c in low):
            hdr = i
            break
    region = grid[:hdr] if hdr < len(grid) else grid[:8]

    # total tally: first row in the region with a 'Vote:' cell + >=3 numbers
    for cells in region:
        for i, c in enumerate(cells):
            if c.lower().endswith("vote:"):
                ns = [int(float(x)) for x in cells[i + 1:] if _is_num(x)]
                if len(ns) >= 3:
                    return result, ns[0], ns[1], ns[2]
    # fallback: first region row that simply has >=3 numbers
    for cells in region:
        ns = [int(float(c)) for c in cells if _is_num(c)]
        if len(ns) >= 3:
            return result, ns[0], ns[1], ns[2]
    return result, None, None, None


def _split_detail(s):
    """'NPRR1308 - To recommend approval…' -> ('NPRR1308', 'To recommend…')."""
    s = re.sub(r"\s+", " ", str(s)).strip()
    if " - " in s:
        left, right = s.split(" - ", 1)
        if _RR_CODE.search(left) or "ballot" in left.lower() or "leadership" in left.lower():
            return left.strip(), right.strip()
    m = _RR_CODE.search(s)
    return (m.group(1) if m else None), s


def extract_ballots(folder, docs):
    """Read every ballot spreadsheet (.xls via xlrd, .xlsx via openpyxl) in the
    meeting folder and return voting outcomes with real tallies. The 'Vote' sheet
    gives the result + committee tally (Yes/No/Abstain); the 'Ballot Details'
    sheet gives the actual item/motion text."""
    outcomes, seen = [], set()
    for f in sorted(docs, key=doc_sort_key):
        low = f.lower()
        if "ballot" not in low or not low.endswith((".xls", ".xlsx")):
            continue
        sheets = _read_sheets(os.path.join(folder, f))
        if not sheets:
            continue
        vote = sheets.get("Vote") or next(iter(sheets.values()), [])
        result, yes, no, ab = _parse_vote_sheet(vote)
        details = [str(r[0]).strip() for r in sheets.get("Ballot Details", [])
                   if r and str(r[0]).strip()]
        combined = "combined" in low or "combo" in low or len(details) >= 4

        def add(item, motion, fr, ag, abst, res):
            motion = re.sub(r"\s+", " ", str(motion)).strip()[:240]
            key = (item, motion[:50])
            if motion and key not in seen:
                seen.add(key)
                outcomes.append({"item": item, "motion": motion, "result": res,
                                 "for": fr, "against": ag, "abstain": abst})

        if combined:
            n = len([d for d in details if d]) or None
            add("Combined Ballot",
                f"Approve the combined ballot{f' ({n} items)' if n else ''}",
                yes, no, ab, result)
            for d in details:                       # list each item (shares result)
                it, mo = _split_detail(d)
                add(it, mo, None, None, None, result)
        elif details:
            for d in details:                       # each item carries this tally
                it, mo = _split_detail(d)
                add(it, mo, yes, no, ab, result)
        else:
            m = _RR_CODE.search(f)
            add(m.group(1) if m else None,
                clean_doc_title(f) or f, yes, no, ab, result)
        if len(outcomes) >= 40:
            break
    return outcomes


# ── Minutes parsing: debates + real voting outcomes ──────────────────────────
NUM = {"one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6,
       "seven": 7, "eight": 8, "nine": 9, "ten": 10}
RR = re.compile(r"\b((?:NPRR|NOGRR|PGRR|RRGRR|SCR|RMGRR|COPMGRR|VCMRR|OBDRR|"
                r"SMOGRR|NGRR|PGRR)\d{2,4})\b")
STRONG = re.compile(r"\b(concern|oppos|object|discuss|withdr|requested additional|"
                    r"further discuss|disagree|debat|reiterat|caution|challenge|"
                    r"friendly amendment|took no action)\w*", re.I)
EXCLUDE = re.compile(r"no objection|there were no objection|antitrust|to order|"
                     r"adjourn|no other business", re.I)
_WORD = {"app": None, "count": 0}  # reused Word instance, recycled periodically
_WORD_RECYCLE = 25                 # restart Word every N opens (avoids degradation)


def _word_paragraphs(path):
    """Read a legacy .doc via a reused Word instance; [] on any failure.

    Reads the whole document in one `Content.Text` call (splitting on the \\r
    paragraph mark) rather than iterating the Paragraphs COM collection, which
    returns incomplete results when many files are opened in sequence. Dialogs
    and macros are suppressed so a stray document can't block the batch.
    """
    try:
        import win32com.client
        # Recycle Word periodically — a long-lived instance starts returning
        # truncated Content.Text after ~150 opens.
        if _WORD["app"] is not None and _WORD["count"] >= _WORD_RECYCLE:
            close_word()
        if _WORD["app"] is None:
            app = win32com.client.Dispatch("Word.Application")
            app.Visible = False
            try:
                app.DisplayAlerts = 0            # wdAlertsNone
                app.AutomationSecurity = 3       # msoAutomationSecurityForceDisable
            except Exception:
                pass
            _WORD["app"] = app
            _WORD["count"] = 0
        _WORD["count"] += 1
        doc = _WORD["app"].Documents.Open(
            os.path.abspath(path), ReadOnly=True, ConfirmConversions=False,
            AddToRecentFiles=False)
        text = doc.Content.Text
        doc.Close(False)
        return text.split("\r")
    except Exception:
        return []


def close_word():
    if _WORD["app"] is not None:
        try:
            _WORD["app"].Quit()
        except Exception:
            pass
        _WORD["app"] = None
        _WORD["count"] = 0


def minutes_paragraphs(path):
    low = path.lower()
    if low.endswith(".docx"):
        try:
            from docx import Document
            return [p.text for p in Document(path).paragraphs]
        except Exception:
            return []
    if low.endswith(".pdf"):
        return _pdf_sentences(path)
    return _word_paragraphs(path)  # .doc


def _pdf_sentences(path):
    """PDF minutes: de-wrap lines and split into sentences so a motion and its
    result land in parseable units for parse_minutes()."""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = "\n".join((p.extract_text() or "") for p in pdf.pages)
    except Exception:
        return []
    text = re.sub(r"-\n", "", text)            # join hyphenated line breaks
    text = re.sub(r"\s*\n\s*", " ", text)       # de-wrap
    for ab in ("Mr.", "Ms.", "Mrs.", "Dr.", "No.", "Sr.", "Jr.", "St.",
               "Inc.", "Co.", "vs.", "U.S.", "a.m.", "p.m."):
        text = text.replace(ab, ab.replace(".", "\x00"))
    sents = re.split(r"(?<=[.;:])\s+(?=[A-Z(])", text)
    return [s.replace("\x00", ".").strip() for s in sents if s.strip()]


def _result_text(low):
    if "fail" in low:
        base = "Failed"
    elif "unanim" in low:
        return "Carried unanimously"
    else:
        base = "Carried"
    opp = re.search(r"(one|two|three|four|five|six|\d+)\s+(?:opposing|objection|negative)", low)
    ab = re.search(r"(one|two|three|four|five|six|\d+)\s+abstention", low)
    extra = []
    if opp:
        extra.append(f"{NUM.get(opp.group(1), opp.group(1))} opposing")
    if ab:
        extra.append(f"{NUM.get(ab.group(1), ab.group(1))} abstaining")
    return base + (" — " + ", ".join(extra) if extra else "")


def parse_minutes(paras):
    """Return (debates, voting_outcomes) parsed from minutes paragraphs."""
    outcomes, debates, seen_o, seen_d = [], [], {}, set()
    last_item = None
    for raw in paras:
        t = re.sub(r"\s+", " ", raw).strip()
        if not t:
            continue
        low = t.lower()
        if RR.search(t) and len(t) < 120:
            last_item = RR.search(t).group(1)
        has_moved = "moved to" in low
        has_result = ("motion carried" in low or "carried with" in low
                      or "carried unanim" in low or "motion fail" in low
                      or "motion passed" in low or "passed with" in low
                      or "passed unanim" in low)
        if has_moved:
            if "combined ballot" in low or "combo ballot" in low:
                item = "Combined Ballot"
            else:
                item = RR.search(t).group(1) if RR.search(t) else last_item
            res = _result_text(low) if has_result else None
            mv = re.search(r"moved to (.+?)(?:\.\s|\.$|$)", t, re.I)
            motion = re.sub(r"\s+", " ", (mv.group(1) if mv else t)).strip()[:240]
            k = (item, motion[:50])
            if k in seen_o:
                if res and not seen_o[k]["result"]:
                    seen_o[k]["result"] = res
                continue
            o = {"item": item, "motion": motion, "result": res,
                 "for": None, "against": None, "abstain": None}
            seen_o[k] = o
            outcomes.append(o)
        elif has_result:  # standalone "The motion carried…" → attach to last open outcome
            res = _result_text(low)
            for o in reversed(outcomes):
                if o["result"] is None:
                    o["result"] = res
                    break
        elif STRONG.search(t) and not EXCLUDE.search(t) and 70 <= len(t) <= 600:
            if low[:40] not in seen_d:
                seen_d.add(low[:40])
                debates.append(t if len(t) <= 320 else t[:317].rstrip() + "…")
    return debates[:8], outcomes[:25]


_MONTHS = (r"jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec")


def clean_doc_title(fname):
    """Turn a document filename into a readable title for the documents-based
    fallback summary, e.g. '9. ROS_Update_to_TAC_04_29_2026.pptx' -> 'ROS Update
    to TAC'."""
    t = os.path.splitext(fname)[0]
    t = re.sub(r"^\s*\d+(?:\.\d+)?[.\-\s]+", "", t)         # leading '04. ' / '5.-' / '9.1 '
    t = re.sub(r"[_\-]+", " ", t)
    t = re.sub(rf"\b(?:{_MONTHS})[a-z]*\.?\s*\d{{0,2}}\b", "", t, flags=re.I)  # 'Apr29'
    t = re.sub(r"\b\d{1,2}[ ./-]\d{1,2}[ ./-]\d{2,4}\b", "", t)  # dates (any sep)
    t = re.sub(r"\b\d{6,8}\b", "", t)                       # 042926 / 20260429
    t = re.sub(r"\b20\d{2}\b", "", t)                       # stray year
    t = re.sub(r"\b(?:rev\s*\d+|v\d+|final|public|draft)\b", "", t, flags=re.I)
    t = re.sub(r"\s+", " ", t).strip(" -—_·.")
    return t


def summarize_from_documents(folder, docs):
    """Fallback when a meeting has no verified (date-matched) minutes: read
    through the other documents in the folder and summarize what went on —
    a few discussion sentences pulled from readable reports/memos, plus the
    titles of the reports, updates, and presentations that were given."""
    activity, seen = [], set()
    readable = []
    for f in docs:
        low = f.lower()
        if low.endswith((".tmp", ".extracted")):
            continue
        if "agenda" in low or "minute" in low or "ballot" in low:
            continue
        if not low.endswith((".pptx", ".ppt", ".pdf", ".docx", ".doc")):
            continue
        title = clean_doc_title(f)
        if title and len(title) > 3 and title.lower() not in seen:
            seen.add(title.lower())
            activity.append(title)
        if low.endswith((".docx", ".pdf")) and any(
                k in low for k in ("report", "memo", "summary", "update",
                                   "comment", "recommend", "overview")):
            readable.append(f)

    discussion, seen_d = [], set()
    for f in readable[:2]:                       # bounded: read at most 2 docs
        for s in minutes_paragraphs(os.path.join(folder, f)):
            s = re.sub(r"\s+", " ", s).strip()
            if 60 <= len(s) <= 320 and STRONG.search(s) and not EXCLUDE.search(s):
                k = s.lower()[:40]
                if k not in seen_d:
                    seen_d.add(k)
                    discussion.append(s if len(s) <= 300 else s[:297] + "…")
            if len(discussion) >= 4:
                break
        if len(discussion) >= 4:
            break

    bullets = discussion + [f"{a} — presented" for a in activity]
    return bullets[:8]


def wg_reports(files):
    found = []
    for f in files:
        low = f.lower()
        if not any(k in low for k in ("report", "update", "to-ros", "to_ros", "to-tac")):
            continue
        for wg in WG_ABBRS:
            if re.search(rf"\b{wg}\b", f, re.I) and wg not in found:
                found.append(wg)
    return found


def meeting_type(files):
    low = " ".join(files).lower()
    if "email" in low and "vote" in low:
        return "Email Vote"
    if "webex" in low:
        return "WebEx"
    return "Regular"


def doc_sort_key(name):
    """Order documents by the leading number before the first '-' (numeric
    ascending, so 2 < 10), e.g. '02.-Agenda…' < '10.-…' < '11.-…'. Decimal
    prefixes like '5.1-' sort within their group; files with no leading number
    sort after the numbered ones, alphabetically."""
    m = re.match(r"\s*(\d+(?:\.\d+)?)", name)
    if m:
        return (0, float(m.group(1)), name.lower())
    return (1, 0.0, name.lower())


def build_profile(abbr, date, folder):
    files = sorted(f for f in os.listdir(folder)
                   if os.path.isfile(os.path.join(folder, f)))
    docs = sorted((f for f in files
                   if not f.lower().endswith((".zip", ".tmp", ".extracted"))
                   and f != "_manifest.json"),
                  key=doc_sort_key)

    # Agenda — prefer .docx (cleanest), else a PDF agenda.
    agenda = []
    agenda_file = next((f for f in docs if "agenda" in f.lower()
                        and f.lower().endswith(".docx")), None)
    if not agenda_file:
        agenda_file = next((f for f in docs if "agenda" in f.lower()
                            and f.lower().endswith(".pdf")), None)
    if agenda_file:
        agenda = extract_agenda(os.path.join(folder, agenda_file))
    # Read every ballot spreadsheet for voting outcomes with real tallies.
    ballots = extract_ballots(folder, docs)

    topics = [a for a in agenda if not TOPIC_SKIP.search(a)][:8]

    # Study THIS meeting's own minutes (folders often also hold the prior
    # meeting's draft minutes — match by date so we summarize the right meeting).
    # Prefer approved over draft, and cleaner formats (.docx > .doc > .pdf).
    debates, mins_outcomes = [], []
    yyyymmdd = date.replace("-", "")
    mmddyyyy = date[5:7] + date[8:10] + date[0:4]
    def is_this_meeting(f):
        dg = re.sub(r"\D", "", f)
        return yyyymmdd in dg or mmddyyyy in dg
    FMT_RANK = {".docx": 3, ".doc": 2, ".pdf": 1}
    mins = [f for f in docs
            if "minute" in f.lower() and f.lower().endswith((".docx", ".doc", ".pdf"))
            and is_this_meeting(f)]
    minutes_file = max(
        mins, key=lambda f: (10 if "approved" in f.lower() else 0)
                            + FMT_RANK[os.path.splitext(f.lower())[1]]) if mins else None
    if minutes_file:
        debates, mins_outcomes = parse_minutes(
            minutes_paragraphs(os.path.join(folder, minutes_file)))
    # No verified minutes (or none with extractable debates) → summarize what
    # went on from the meeting's other documents.
    if not debates:
        debates = summarize_from_documents(folder, docs)

    # meeting_summary (meeting-level) — distinct from group_summary (group-level).
    # Prefer ballot-derived outcomes (real Yes/No/Abstain tallies + Passed/Failed),
    # then minutes-derived outcomes when there are no ballot spreadsheets.
    voting_outcomes = ballots if ballots else mins_outcomes
    meeting_sum = {"topics": topics, "debates": debates,
                   "voting_outcomes": voting_outcomes}

    full, chair, vc, gsum = group_summary(abbr)
    mmddyyyy = date[5:7] + date[8:10] + date[0:4]
    slug = CAL_SLUG.get(abbr, f"{abbr}-Meeting")

    return {
        "committee": abbr,
        "committee_full_name": full,
        "meeting_date": date,
        "meeting_type": meeting_type(files),
        "calendar_url": f"https://www.ercot.com/calendar/{mmddyyyy}-{slug}",
        "chair": chair,
        "vice_chair": vc,
        "group_summary": gsum,
        "meeting_summary": meeting_sum,
        "agenda_items": agenda,
        "ballot_results": ballots,
        "working_group_reports": wg_reports(docs),
        "key_discussion_topics": topics,
        "action_items": [],
        "documents": docs,
        "next_meeting_date": None,
    }


def main():
    overwrite = "--overwrite" in sys.argv
    only = [a for a in sys.argv[1:] if not a.startswith("--")]
    written = skipped = errors = 0
    try:
        for abbr in sorted(os.listdir(ROOT)):
            if only and abbr not in only:
                continue
            cdir = os.path.join(ROOT, abbr)
            if not os.path.isdir(cdir):
                continue
            for date in sorted(d for d in os.listdir(cdir) if DATE_RE.match(d)):
                folder = os.path.join(cdir, date)
                if not os.path.isdir(folder):
                    continue
                qr = os.path.join(folder, "Quick runs")
                out = os.path.join(qr, f"{abbr}-{date} Profile.json")
                if os.path.exists(out) and not overwrite:
                    skipped += 1; continue
                try:
                    prof = build_profile(abbr, date, folder)
                    os.makedirs(qr, exist_ok=True)
                    with open(out, "w", encoding="utf-8") as fh:
                        json.dump(prof, fh, indent=2, ensure_ascii=False)
                    written += 1
                    if written % 100 == 0:
                        print(f"  ... {written} written")
                except Exception as e:
                    errors += 1
                    print(f"  ERROR {abbr}/{date}: {e}")
    finally:
        close_word()
    print(f"\nDone. written={written} skipped(existing)={skipped} errors={errors}")


if __name__ == "__main__":
    main()
